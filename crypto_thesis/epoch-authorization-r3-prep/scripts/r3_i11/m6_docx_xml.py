# -*- coding: utf-8 -*-
"""M6: dump one equation paragraph XML for inspection."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
OUT = ROOT / "docs/midterm-report/m6/output/王威-专业学位研究生学位论文中期考评表-M6候选稿.docx"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(OUT))
    cell = doc.tables[0].rows[4].cells[0]
    shown = 0
    for p in cell.paragraphs:
        omats = p._p.findall(f".//{M}oMath")
        if omats and re.match(r"^\(\d+\)", p.text.strip()):
            print("=== paragraph text:", repr(p.text[:60]))
            for om in omats:
                parent = om.getparent()
                print("parent tag:", etree.QName(parent).localname)
                print("om tag:", etree.QName(om).localname)
                print("om xml tail (600):", etree.tostring(om, encoding="unicode")[:600])
                print()
            shown += 1
            if shown >= 2:
                break


if __name__ == "__main__":
    main()
