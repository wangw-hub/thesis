"""Verify M2 DOCX: cover, tables, cells, images."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
DOCX = ROOT / "docs/midterm-report/m2/output/王威-专业学位研究生学位论文中期考评表-M2候选稿.docx"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    d = Document(str(DOCX))
    print("tables:", len(d.tables), "| inline shapes:", len(d.inline_shapes))
    cover = "\n".join(p.text for p in d.paragraphs[:16])
    for kw in ["202422081113", "王威", "高建彬", "王鹏", "计算机技术", "计算机科学与工程学院"]:
        print(f"cover {kw}:", kw in cover)
    t0 = d.tables[0]
    print("t0 rows:", len(t0.rows))
    r4 = t0.rows[4].cells[0].text
    r6 = t0.rows[6].cells[0].text
    print("progress len:", len(r4), "| has 15120:", "15120" in r4, "| has 图1:", "图1" in r4)
    print("stage results len:", len(r6), "| has 拟投稿:", "拟投稿" in r6, "| head:", r6[:60].replace(chr(10), " "))
    t1 = d.tables[1]
    p_len = len(t1.rows[0].cells[0].text)
    s_len = len(t1.rows[1].cells[0].text)
    print("problems len:", p_len, "| solutions len:", s_len)
    for kw in ["总体来看", "（1）论文整体学术逻辑", "阶段3（2027年1月"]:
        print(f"  {kw}:", kw in t1.rows[0].cells[0].text + t1.rows[1].cells[0].text)


if __name__ == "__main__":
    main()
