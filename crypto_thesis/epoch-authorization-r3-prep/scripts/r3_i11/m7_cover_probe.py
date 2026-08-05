# -*- coding: utf-8 -*-
"""M7: inspect template/M6 cover fields (date, checks, blank page)."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
TEMPLATE = Path(r"D:\Users\wangw\Desktop\中期和小论文\王威专业学位研究生学位论文中期考评表.docx")
M6_DOCX = ROOT / "docs/midterm-report/m6/output/王威-专业学位研究生学位论文中期考评表-M6候选稿.docx"


def dump_cover(path: Path, label: str) -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print("=" * 20, label)
    doc = Document(str(path))
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"p{i:02d}: {p.text[:110]}")
    for ti, t in enumerate(doc.tables):
        if ti > 1:
            break
        print(f"--- table {ti} ---")
        for row in t.rows:
            cells = [c.text.replace("\n", " ")[:60] for c in row.cells]
            print(" | ".join(cells))


def main() -> None:
    dump_cover(TEMPLATE, "TEMPLATE")
    dump_cover(M6_DOCX, "M6 DOCX")


if __name__ == "__main__":
    main()
