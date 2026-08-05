# -*- coding: utf-8 -*-
"""I17: rebuild the thesis DOCX (V2) on the official UESTC template basis.

Uses the official cover/flyleaf template as the base document, fills its
fields, appends the school's declaration pages, then builds the front matter
(abstract/abstract/toc), the body chapters with official styles, headers,
footers, page numbering, tables, figures, equations and references.
"""
from __future__ import annotations

import copy
import io
import json
import os
import re
import shutil
import stat
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
COVER_TEMPLATE = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\电子科技大学研究生学位论文封面及扉页 - 适用于专业学位硕士_081705087525.docx")
SPEC_TEMPLATE = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\电子科技大学研究生学位论文撰写规范- 适用于中国学生 - 副本_031543351520.docx")
SRC = ROOT / "docs/final-manuscript/i17/I17-SOURCE.md"
OUT = ROOT / "docs/final-manuscript/output"
OUT_DOCX = OUT / "THESIS-FORMAT-CANDIDATE-V2.docx"
RC2_FIG = Path(r"D:\Research\crypto_thesis\epoch-authorization\docs\thesis-drafts\research-content-2-final\figures")
RC3_FIG = ROOT / "experiments/r3/formal/figures/i12-final"
RC3_TAB = ROOT / "experiments/r3/formal/tables/i12-final"
I12_RQ = ROOT / "docs/research-content-3-implementation/i12/formal-rq-results.json"
MML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"

TITLE = "面向非连续时间约束的区块链数据共享关键技术研究及实现"
EN_TITLE = "Research and Implementation of Key Technologies for Blockchain Data Sharing under Non-Continuous Time Constraints"
STUDENT_ID = "202422081113"
STUDENT_NAME = "王威"
STUDENT_NAME_EN = "Wang Wei"
COLLEGE = "计算机科学与工程学院（网络空间安全学院）"
COLLEGE_EN = "School of Computer Science and Engineering(School of Cyber Security)"
CATEGORY = "计算机技术"
CATEGORY_EN = "Computer Technology"
SUPERVISOR = "高建彬"
SUPERVISOR_EN = "Gao Jianbin"
ACCESS_DATE = "2026-08-02"


# ---------------------------------------------------------------------------
# OMML
# ---------------------------------------------------------------------------

_XSLT = None


def _transformer():
    global _XSLT
    if _XSLT is None:
        _XSLT = etree.XSLT(etree.parse(MML_XSL))
    return _XSLT


def latex_to_omml(latex: str) -> etree._Element:
    mathml = latex_to_mathml(latex, display="block")
    tree = etree.fromstring(mathml.encode("utf-8"))
    res = _transformer()(tree)
    root = res.getroot()
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    om = root.find(".//m:oMath", ns)
    if om is None:
        om = root
    return copy.deepcopy(om)


def add_omml(p, latex: str) -> None:
    p._p.append(latex_to_omml(latex))


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\s][^*]*?\*|`[^`]+?`|\\\(.*?\\\)|\\\[.*?\\\]|\[\d+(?:,\s*\d+)*\]|\[\d+\])", re.S)


def set_run(run, bold=False, italic=False, code=False, sup=False, east="宋体", latin="Times New Roman", size=None):
    if code:
        latin = "Consolas"
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), east)
    run.bold = bold
    run.italic = italic
    if sup:
        run.font.superscript = True
    if size:
        run.font.size = Pt(size)


def render_inline(p, text: str, size: float = 12, east="宋体"):
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); set_run(r, bold=True, size=size, east=east)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); set_run(r, code=True, size=size - 1, east=east)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = p.add_run(tok[1:-1]); set_run(r, italic=True, size=size, east=east)
        elif tok.startswith("\\(") and tok.endswith("\\)"):
            add_omml(p, tok[2:-2])
        elif tok.startswith("\\[") and tok.endswith("\\]"):
            add_omml(p, tok[2:-2])
        elif re.fullmatch(r"\[\d+(?:,\s*\d+)*\]", tok):
            r = p.add_run(tok); set_run(r, sup=True, size=size, east=east)
        else:
            r = p.add_run(tok); set_run(r, size=size, east=east)


def para_props(p, align=None, first_chars=0, hanging_chars=0, left_cm=0.0, keep_next=False,
               line_pt=20, space_before=0, space_after=0):
    pf = p.paragraph_format
    pf.line_spacing = Pt(line_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if first_chars:
        ind.set(qn("w:firstLineChars"), str(first_chars * 100))
        ind.set(qn("w:firstLine"), str(int(first_chars * 240)))
    if hanging_chars:
        ind.set(qn("w:hangingChars"), str(hanging_chars * 100))
        ind.set(qn("w:hanging"), str(int(hanging_chars * 240)))
    if left_cm:
        ind.set(qn("w:left"), str(int(left_cm * 567)))
    if keep_next:
        pf.keep_with_next = True
    return p


# ---------------------------------------------------------------------------
# Official styles
# ---------------------------------------------------------------------------


def style_fonts(style, east, size, bold=None, latin="Times New Roman"):
    style.font.name = latin
    rpr = style.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), east)
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def setup_styles(doc: Document) -> None:
    heads = [
        ("Heading 1", HEI := "黑体", 15, None, 24, 18, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", "黑体", 14, None, 18, 6, None),
        ("Heading 3", "黑体", 14, None, 12, 6, None),
        ("Heading 4", "黑体", 12, True, 12, 6, None),
    ]
    for name, east, size, bold, before, after, align in heads:
        st = doc.styles[name]
        style_fonts(st, east, size, bold)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = Pt(20)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        st.paragraph_format.keep_with_next = True
        if align is not None:
            st.paragraph_format.alignment = align
        # remove any automatic numbering inherited from the template's heading styles
        ppr = st.element.get_or_add_pPr()
        numpr = ppr.find(qn("w:numPr"))
        if numpr is not None:
            ppr.remove(numpr)
    normal = doc.styles["Normal"]
    style_fonts(normal, "宋体", 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    for name, east, size, bold, before, after in [
        ("FrontHeading", "黑体", 15, None, 0, 18),
        ("图题", "宋体", 10.5, None, 6, 12),
        ("表题", "宋体", 10.5, None, 12, 6),
        ("公式", "宋体", 12, None, 6, 6),
        ("参考文献", "宋体", 10.5, None, 0, 0),
        ("算法题注", "宋体", 10.5, None, 6, 6),
    ]:
        if name not in doc.styles:
            doc.styles.add_style(name, 1)
        st = doc.styles[name]
        style_fonts(st, east, size, bold)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = Pt(20)
        st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    doc.styles["图题"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["表题"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["FrontHeading"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # remove outline level from FrontHeading so TOC excludes it
    st = doc.styles["FrontHeading"]
    ppr = st.element.get_or_add_pPr()
    ol = ppr.find(qn("w:outlineLvl"))
    if ol is None:
        ol = OxmlElement("w:outlineLvl")
        ppr.append(ol)
    ol.set(qn("w:val"), "9")
    ref_style = doc.styles["参考文献"]
    ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ppr = ref_style.element.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:hangingChars"), "200")
    ind.set(qn("w:hanging"), "240")
    ppr.append(ind)


# ---------------------------------------------------------------------------
# Cover filling
# ---------------------------------------------------------------------------


def set_cell_text(cell, text: str, size: float | None = None, east: str | None = None, bold=None):
    p = cell.paragraphs[0]
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    if p.runs:
        r = p.runs[0]
        r.text = text
        if size:
            r.font.size = Pt(size)
        if east:
            rpr = r._element.get_or_add_rPr()
            rf = rpr.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts")
                rpr.insert(0, rf)
            rf.set(qn("w:eastAsia"), east)
        if bold is not None:
            r.bold = bold
    else:
        r = p.add_run(text)
        set_run(r, size=size or 16, east=east or "宋体", bold=bold if bold is not None else True)


def fill_cover(doc: Document) -> None:
    replaces = [
        ("5G移动通信基站天线关键技术及其", "面向非连续时间约束的区块链数据共享"),
        ("特征模分析方法研究", "关键技术研究及实现"),
        ("电子信息", "计算机技术"),
        ("2017XXXXXXXX", STUDENT_ID),
        ("张\u3000某", "王\u3000威"),
        ("Zhang Mou", STUDENT_NAME_EN),
        ("李某某\u3000\u3000教\u3000授", SUPERVISOR),
        ("李某某", SUPERVISOR),
        ("教\u3000授", ""),
        ("电子科学与工程学院", COLLEGE),
        ("TN828.6", "[待填写]"),
        ("621.39", "[待填写]"),
        ("Prof. Li Moumou", SUPERVISOR_EN),
        ("Electromagnetic Field and", CATEGORY_EN),
        ("Microwave Technology", ""),
        ("School of Electronic Science and Engineering", COLLEGE_EN),
        ("Key Technologies and Characteristic Mode", EN_TITLE.split(" under ")[0]),
        ("Analysis Methods for 5G Base Station Antennas", "under Non-Continuous Time Constraints"),
        ("章某某", "[待填写]"),
        ("赵某、王某某、李某、刘某、戴某某", "[待填写]"),
        ("通信工程（此行根据实际情况保留或删除）", ""),
    ]
    rows_to_delete = []
    for tbl in doc.tables:
        for tc in tbl._tbl.iter(qn("w:tc")):
            texts = [t.text or "" for t in tc.iter(qn("w:t"))]
            txt = "".join(texts)
            for old, new in replaces:
                if old and old in txt:
                    set_tc_text(tc, txt.replace(old, new))
                    if old == "通信工程（此行根据实际情况保留或删除）":
                        tr = tc.getparent()
                        if tr is not None and tr.tag == qn("w:tr"):
                            rows_to_delete.append(tr)
                    break
    for tr in rows_to_delete:
        tr.getparent().remove(tr)


def set_tc_text(tc, text: str) -> None:
    """Replace a table cell's content with plain text, keeping the first run's rPr."""
    paras = tc.findall(qn("w:p"))
    if not paras:
        paras = [OxmlElement("w:p")]
        tc.append(paras[0])
    p0 = paras[0]
    # find first run to copy its properties
    rpr_template = None
    for r in p0.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            rpr_template = rpr
        p0.remove(r)
    for p in paras[1:]:
        tc.remove(p)
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(copy.deepcopy(rpr_template))
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p0.append(r)


def add_declaration_pages(doc: Document) -> None:
    spec = Document(str(SPEC_TEMPLATE))
    texts = {}
    for p in spec.paragraphs:
        if p.text.startswith("本人声明所呈交"):
            texts["declaration"] = p.text
        if p.text.startswith("本学位论文作者完全了解"):
            texts["authorization"] = p.text
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("独创性声明"); set_run(r, size=18, east="华文中宋", bold=True)
    para_props(h, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, space_before=24, space_after=18)
    body = doc.add_paragraph(); para_props(body, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    render_inline(body, texts.get("declaration", ""), size=14)
    sig = doc.add_paragraph(); para_props(sig, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=24)
    render_inline(sig, "作者签名：\t\t日期：\u3000\u3000年\u3000\u3000月\u3000\u3000日", size=14)
    doc.add_page_break()
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("论文使用授权"); set_run(r, size=18, east="华文中宋", bold=True)
    para_props(h, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, space_before=24, space_after=18)
    body = doc.add_paragraph(); para_props(body, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    render_inline(body, texts.get("authorization", ""), size=14)
    body = doc.add_paragraph(); para_props(body, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    render_inline(body, "（涉密的学位论文须按照国家及学校相关规定管理，在解密后适用于本授权。）", size=14)
    sig = doc.add_paragraph(); para_props(sig, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=24)
    render_inline(sig, "作者签名：\t\t导师签名：", size=14)
    sig = doc.add_paragraph(); para_props(sig, align=WD_ALIGN_PARAGRAPH.RIGHT)
    render_inline(sig, "日期：\u3000\u3000年\u3000\u3000月\u3000\u3000日", size=14)


# ---------------------------------------------------------------------------
# Fields / headers / footers
# ---------------------------------------------------------------------------


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()

    def el(tag, attrs=None, text=None):
        e = OxmlElement(tag)
        if attrs:
            for k, v in attrs.items():
                e.set(qn(k), v)
        if text:
            e.text = text
        return e

    r1 = OxmlElement("w:r"); r1.append(el("w:fldChar", {"w:fldCharType": "begin"}))
    r2 = OxmlElement("w:r"); r2.append(el("w:instrText", {"xml:space": "preserve"}, ' TOC \\o "1-3" \\h \\z \\u '))
    r3 = OxmlElement("w:r"); r3.append(el("w:fldChar", {"w:fldCharType": "separate"}))
    r4 = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "（目录）"; r4.append(t)
    r5 = OxmlElement("w:r"); r5.append(el("w:fldChar", {"w:fldCharType": "end"}))
    for r in (r1, r2, r3, r4, r5):
        p._p.append(r)


def styleref_field(p, style_name: str) -> None:
    def el(tag, attrs=None, text=None):
        e = OxmlElement(tag)
        if attrs:
            for k, v in attrs.items():
                e.set(qn(k), v)
        if text:
            e.text = text
        return e

    r1 = OxmlElement("w:r"); r1.append(el("w:fldChar", {"w:fldCharType": "begin"}))
    r2 = OxmlElement("w:r"); r2.append(el("w:instrText", {"xml:space": "preserve"}, f' STYLEREF "{style_name}" \\h '))
    r3 = OxmlElement("w:r"); r3.append(el("w:fldChar", {"w:fldCharType": "separate"}))
    r4 = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = style_name; r4.append(t)
    r5 = OxmlElement("w:r"); r5.append(el("w:fldChar", {"w:fldCharType": "end"}))
    for r in (r1, r2, r3, r4, r5):
        p._p.append(r)


def page_field(p) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def set_section_pages(section, fmt: str, start: int, header_text: str | None = None,
                      even_text: str | None = None, styleref: str | None = None) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(2.0)
    section.footer_distance = Cm(2.0)
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))
    # header
    if header_text is not None:
        section.header.is_linked_to_previous = False
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.line_spacing = Pt(20)
        hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        hp.paragraph_format.space_after = Pt(0)
        if styleref:
            styleref_field(hp, styleref)
        else:
            r = hp.add_run(header_text)
            set_run(r, size=10.5)
        # header bottom border (0.75pt single line)
        ppr = hp._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pbdr.append(bottom)
        ppr.append(pbdr)
    if even_text is not None:
        section.even_page_header.is_linked_to_previous = False
        hp = section.even_page_header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(even_text)
        set_run(r, size=10.5)
        ppr = hp._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pbdr.append(bottom)
        ppr.append(pbdr)
    # footer
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_field(fp)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    uf = settings.find(qn("w:updateFields"))
    if uf is None:
        uf = OxmlElement("w:updateFields")
        settings.append(uf)
    uf.set(qn("w:val"), "true")


# ---------------------------------------------------------------------------
# Tables / figures
# ---------------------------------------------------------------------------


def three_line_borders(table, header_rows=1):
    tbl = table._tbl
    tblPr = tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, sz, val in (("top", 12, "single"), ("bottom", 12, "single"),
                          ("left", 0, "none"), ("right", 0, "none"),
                          ("insideH", 0, "none"), ("insideV", 0, "none")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)
    for cell in table.rows[header_rows - 1].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "000000")
        tcB.append(bottom)
        tcPr.append(tcB)


def add_md_table(doc, rows, caption=None):
    rows = [r for r in rows if any(str(c).strip() for c in r)]
    if not rows:
        return None
    if caption:
        p = doc.add_paragraph(style="表题")
        render_inline(p, caption, size=10.5)
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_cm = 15.0
    col_w = width_cm / ncols
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.width = Cm(col_w)
            txt = row[ci] if ci < len(row) else ""
            p0 = cell.paragraphs[0]
            p0.paragraph_format.line_spacing = Pt(14)
            p0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p0.paragraph_format.space_after = Pt(2)
            p0.paragraph_format.space_before = Pt(2)
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            render_inline(p0, txt, size=10.5)
            if ri < 1:
                for r in p0.runs:
                    r.bold = True
    three_line_borders(table)
    return table


def rc3_table_rows(name: str) -> list[list[str]]:
    if name == "表6-1":
        d = json.load(io.open(RC3_TAB / "table-run-flow-eligibility.json", encoding="utf-8"))
        rows = d["rows"]
        by_exp = {}
        for r in rows:
            by_exp.setdefault(r["experiment"], []).append(r)
        out = [["实验", "配置数", "有效运行", "VALID_SUCCESS", "VALID_EXPECTED_FAIL_CLOSED"]]
        for exp in sorted(by_exp):
            rs = by_exp[exp]
            out.append([exp, str(len({r["config"] for r in rs})),
                        str(sum(1 for r in rs if r.get("valid"))),
                        str(sum(1 for r in rs if r.get("disposition") == "VALID_SUCCESS")),
                        str(sum(1 for r in rs if r.get("disposition") == "VALID_EXPECTED_FAIL_CLOSED"))])
        out.append(["合计", "-", str(len(rows)),
                    str(sum(1 for r in rows if r.get("disposition") == "VALID_SUCCESS")),
                    str(sum(1 for r in rows if r.get("disposition") == "VALID_EXPECTED_FAIL_CLOSED"))])
        return out
    if name == "表6-2":
        d = json.load(io.open(I12_RQ, encoding="utf-8"))
        out = [["配置", "n", "中位数/ms", "IQR/ms", "均值/ms", "95% Bootstrap CI/ms"]]
        for rq in ("RQ-2", "RQ-3"):
            for label in sorted(d["cards"][rq].get("levels", {})):
                v = d["cards"][rq]["levels"][label]
                ci = v.get("ci95") or [None, None]
                ci_txt = f"[{ci[0]:.1f}, {ci[1]:.1f}]" if ci[0] is not None else "-"
                out.append([label.replace("recipient=", "接收者").replace("affected=", "资源").replace("body=", "Body "),
                            str(v["n"]), f"{v['median']:.1f}", f"{v['iqr']:.1f}", f"{v['mean']:.1f}", ci_txt])
        return out
    if name == "表6-3":
        d = json.load(io.open(RC3_TAB / "table-matched-local-kubo-recovery.json", encoding="utf-8"))
        cells = d["cells"]
        out = [["故障", "对象来源", "有效(n)", "恢复判定", "修复动作", "时长中位数/ms"]]
        for f in ["NONE", "CORRUPT_RESTORE", "CID_MISMATCH", "BOTH_MISSING"]:
            for src in ("LOCAL_ONLY", "KUBO_REPLICA"):
                c = cells.get(f, {}).get(src)
                if not c:
                    continue
                rec = "/".join(f"{k}:{v}" for k, v in (c.get("recoveryDispositions") or {}).items()) or "-"
                rep = "/".join(f"{k}:{v}" for k, v in (c.get("repairActions") or {}).items()) or "-"
                med = c.get("durationMedianMs")
                out.append([f, src, str(c["n"]), rec, rep, f"{med:.1f}" if med is not None else "-"])
        return out
    if name == "表6-4":
        d = json.load(io.open(RC3_TAB / "table-release-decision-outcome.json", encoding="utf-8"))
        dec = d["decisions"]
        return [["释放判定", "运行数"],
                ["ALLOWED_AFTER_CURRENT_HEADER_ONLY", str(dec.get("ALLOWED_AFTER_CURRENT_HEADER_ONLY", 0))],
                ["DENIED", str(dec.get("DENIED", 0))], ["错误材料释放", str(d.get("wrongMaterialRelease", 0))]]
    if name == "表6-5":
        d = json.load(io.open(RC3_TAB / "table-environment-fingerprint.json", encoding="utf-8"))
        f = d.get("fingerprint", d)
        rows = [["项目", "值"]]
        pairs = [
            ("主机/角色", f"{f.get('host')}/{f.get('role')}"),
            ("CPU", f.get("cpuModel", "")),
            ("物理/逻辑核", f"{f.get('physicalCores')}/{f.get('logicalCores')}"),
            ("内存", f"{f.get('ramBytes', 0) / 2**30:.2f} GiB"),
            ("OS/内核", f"{f.get('os')} {f.get('kernel')}"),
            ("虚拟化/网络", f"{f.get('virtualization')}；{f.get('network')}"),
            ("Python", f.get("pythonVersion", "")),
            ("Besu", f.get("besuVersion", "")),
            ("PostgreSQL", f.get("postgresqlVersion", "")),
            ("Kubo", f.get("kuboVersion", "")),
        ]
        for k, v in pairs:
            rows.append([k, str(v)])
        return rows
    return []


RC3_FIGS = {"图6-1": "fig-rq2-header-only-duration.png",
            "图6-2": "fig-rq3-body-rotation-duration.png",
            "图6-3": "fig-rq5-recovery-local-kubo.png"}


def resolve_image(path: str) -> Path:
    p = Path(path.replace("/", "\\"))
    cands = []
    if p.is_absolute():
        cands.append(p)
    else:
        cands += [ROOT / p, ROOT.parent / p, RC2_FIG / p.name]
    for c in cands:
        if c.is_file():
            return c
    return p


# ---------------------------------------------------------------------------
# Reference reformat (school GB/T 7714 details)
# ---------------------------------------------------------------------------

ACRONYMS = ["TRBAC", "JCS", "JWT", "HPKE", "EdDSA", "OAuth", "JSON", "XML", "DHT", "P2P",
            "RFC", "SHA-256", "Ed25519", "AES", "QuickCheck", "Haskell", "Bootstrap",
            "Jackknife", "ICDE", "ICFP", "ACM", "IEEE", "IPFS", "Edwards"]


def sentence_title(title: str) -> str:
    t = title[0].upper() + title[1:].lower()
    for a in sorted(ACRONYMS, key=len, reverse=True):
        t = re.sub(rf"(?<![A-Za-z]){re.escape(a)}(?![A-Za-z])", a, t, flags=re.I)
    return t


def fix_authors(authors: str) -> str:
    parts = [a.strip() for a in authors.split(",") if a.strip()]
    if len(parts) > 3:
        parts = parts[:3] + ["et al."]
    return ", ".join(parts)


def reformat_ref(entry: str) -> str:
    m = re.match(r"^(\[\d+\])\s+(.*)$", entry.strip())
    if not m:
        return entry
    num, body = m.group(1), m.group(2)
    body = body.replace("[C]//", "[C]. ")
    body = body.replace("[C]//", "[C]. ")
    # authors. title[type]. rest
    m2 = re.match(r"^(.*?[A-Za-z\u4e00-\u9fff]\.)\s+(.*?)(\[[A-Z/]+\]\.)\s*(.*)$", body, re.S)
    if not m2:
        return entry
    authors, title, typ, rest = m2.group(1), m2.group(2), m2.group(3), m2.group(4)
    authors = fix_authors(authors)
    title = sentence_title(title) if re.search(r"[A-Za-z]", title) else title
    rest = re.sub(r"\s*DOI:\s*10\.[0-9A-Za-z./-]+\.?\s*$", "", rest)
    return f"{num} {authors} {title}{typ} {rest}"


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------


def chapter_key(t: str) -> str | None:
    m = re.match(r"^(第[一二三四五六七1-7]章|参考文献|附录A|致谢|攻读学位期间)", t)
    return m.group(1) if m else None


def heading_level(t: str) -> int | None:
    m = re.match(r"^(\d+)\.(\d+)(?:\.(\d+))?\s", t)
    if m:
        return 2 if not m.group(3) else 3
    return None


def parse_blocks(text: str) -> list[dict]:
    lines = text.splitlines()
    blocks: list[dict] = []
    i = 0
    n = len(lines)
    while i < n:
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", s)
        if m:
            blocks.append({"type": "image", "alt": m.group(1), "path": m.group(2)})
            i += 1
            continue
        if s.startswith("```"):
            lang = s[3:].strip()
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(buf)})
            continue
        if s.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            data = [r for r in rows if not all(re.fullmatch(r":?-{3,}:?", c) for c in r)]
            data = [r for r in data if any(c for c in r)]
            if data:
                blocks.append({"type": "table", "rows": data})
            else:
                blocks.append({"type": "para", "text": lines[i - 1].strip()})
            continue
        if s.startswith("\\["):
            buf = [lines[i]]
            i += 1
            while i < n and "\\]" not in lines[i]:
                buf.append(lines[i])
                i += 1
            if i < n:
                buf.append(lines[i])
                i += 1
            content = "\n".join(buf).strip()
            content = re.sub(r"^\\\[\s*", "", content)
            content = re.sub(r"\s*\\\]$", "", content)
            blocks.append({"type": "eq", "latex": content})
            continue
        if re.match(r"^\*\*(图|表|算法)", s):
            blocks.append({"type": "caption", "text": s})
            i += 1
            continue
        blocks.append({"type": "para", "text": lines[i].strip()})
        i += 1
    return blocks


UNCAPTIONED_TITLES = {
    ("4", "字段"): "NTP1 固定宽度大端编码字段",
    ("4", "阶段"): "各阶段复杂度分析",
    ("4", "场景"): "不同策略场景下三种表示的适用性",
    ("5", "目标"): "安全目标、机制与证据及结论边界",
}


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    shutil.copy(COVER_TEMPLATE, OUT_DOCX)
    os.chmod(OUT_DOCX, stat.S_IWRITE | stat.S_IREAD)
    doc = Document(str(OUT_DOCX))
    setup_styles(doc)
    sec0 = doc.sections[0]
    sec0.page_width = Cm(21.0)
    sec0.page_height = Cm(29.7)
    sec0.top_margin = Cm(3.0)
    sec0.bottom_margin = Cm(3.0)
    sec0.left_margin = Cm(3.0)
    sec0.right_margin = Cm(3.0)
    sec0.header_distance = Cm(2.0)
    sec0.footer_distance = Cm(2.0)
    sec0.footer.is_linked_to_previous = False
    for p in sec0.footer.paragraphs:
        p.text = ""
    fill_cover(doc)
    add_declaration_pages(doc)

    text = io.open(SRC, encoding="utf-8").read()
    blocks = parse_blocks(text)

    # front matter collection
    front: list[dict] = []
    body_blocks: list[dict] = []
    state = "front"
    for b in blocks:
        if b["type"] == "heading":
            t = b["text"]
            if t in ("中文摘要", "关键词", "Abstract", "Keywords", "目录"):
                state = "front"
            elif chapter_key(t):
                state = "body"
        (front if state == "front" else body_blocks).append(b)

    # build front matter section
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_pages(sec2, fmt="upperRoman", start=1, header_text=None, styleref="FrontHeading")
    kw_zh = "非连续时间约束；区块链数据共享；可信授权执行；版本化密文；前瞻性撤销"
    kw_en = "Non-Continuous Time Constraint; Blockchain Data Sharing; Trusted Authorization; Versioned Ciphertext; Forward-Looking Revocation"
    mode = None
    for b in front:
        if b["type"] == "heading":
            t = b["text"]
            if t == "中文摘要":
                mode = "abs"; p = doc.add_paragraph(style="FrontHeading"); p.add_run("摘\u3000要")
            elif t == "关键词":
                mode = "kw"
            elif t == "Abstract":
                mode = "abs"; p = doc.add_paragraph(style="FrontHeading"); p.add_run("ABSTRACT")
                p.paragraph_format.page_break_before = True
            elif t == "Keywords":
                mode = "kw"
            elif t == "目录":
                mode = "toc"; p = doc.add_paragraph(style="FrontHeading"); p.add_run("目\u3000录")
                p.paragraph_format.page_break_before = True
                add_toc(doc)
            continue
        if b["type"] == "para":
            s = b["text"].strip()
            if mode == "kw":
                p = doc.add_paragraph()
                para_props(p, first_chars=0)
                label = "Keywords: " if s[:1].isascii() else "关键词："
                r = p.add_run(label)
                set_run(r, bold=True, size=12)
                render_inline(p, kw_zh if s.startswith("非连续") else kw_en, size=12)
                mode = None
            elif mode == "abs":
                p = doc.add_paragraph()
                para_props(p, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                render_inline(p, s, size=12)

    # build body section
    sec3 = doc.add_section(WD_SECTION.NEW_PAGE)
    sectPr = sec3._sectPr
    eh = OxmlElement("w:evenAndOddHeaders")
    sectPr.append(eh)
    set_section_pages(sec3, fmt="decimal", start=1, header_text=None, even_text="电子科技大学硕士学位论文", styleref="Heading 1")

    stats = {"paras": 0, "tables": 0, "figures": 0, "equations": 0, "refs": 0,
             "algorithms": 0, "headings": 0, "captions": 0, "manualBreaks": 0}
    eq_counters: dict[str, int] = {}
    tbl_counters: dict[str, int] = {}
    current_chapter = "前置"
    in_refs = False
    xiezhi_done = False
    chengguo_done = False
    i = 0
    while i < len(body_blocks):
        b = body_blocks[i]
        bt = b["type"]
        if bt == "heading":
            t = b["text"]
            key = chapter_key(t)
            if key:
                if key == "参考文献":
                    if not xiezhi_done:
                        add_xiezhi(doc, stats)
                        xiezhi_done = True
                    in_refs = True
                p = doc.add_paragraph(style="Heading 1")
                p.paragraph_format.page_break_before = True
                p.add_run(t)
                current_chapter = t
                stats["headings"] += 1
                i += 1
                continue
            lvl = heading_level(t)
            if lvl is None:
                lvl = min(b["level"], 4)
            p = doc.add_paragraph(style=f"Heading {lvl}")
            p.add_run(t)
            stats["headings"] += 1
            i += 1
            continue
        if bt == "image":
            path = resolve_image(b["path"])
            alt = b["alt"].strip()
            p = doc.add_paragraph()
            para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line_pt=14, space_before=6)
            run = p.add_run()
            try:
                run.add_picture(str(path), width=Cm(13.5))
            except Exception:
                run.add_text("[图片缺失]")
            cap = doc.add_paragraph(style="图题")
            render_inline(cap, alt, size=10.5)
            stats["figures"] += 1
            stats["captions"] += 1
            if i + 1 < len(body_blocks) and body_blocks[i + 1]["type"] == "para":
                nxt = body_blocks[i + 1]["text"].strip()
                if nxt in (alt, alt + "。"):
                    i += 1
            i += 1
            continue
        if bt == "caption":
            raw = b["text"]
            m_rc3 = re.match(r"^\*\*(图|表)\s*(\d+-\d+)([^*]*?)\*\*(.*)$", raw)
            if m_rc3 and m_rc3.group(2).startswith("6-"):
                kind, num, inner, rest = m_rc3.groups()
                cap_txt = f"{kind}{num}{inner} {rest}".strip()
                if kind == "图":
                    fname = RC3_FIGS.get(f"图{num}")
                    if fname:
                        p = doc.add_paragraph()
                        para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line_pt=14, space_before=6)
                        run = p.add_run()
                        run.add_picture(str(RC3_FIG / fname), width=Cm(13.5))
                        stats["figures"] += 1
                    cap = doc.add_paragraph(style="图题")
                    clean = re.sub(r"\s+", " ", cap_txt.replace("（E2）", "").replace("（E3）", "").replace("（E5）", "")).strip()
                    render_inline(cap, clean, size=10.5)
                    stats["captions"] += 1
                    if i + 1 < len(body_blocks) and body_blocks[i + 1]["type"] == "para" and \
                            body_blocks[i + 1]["text"].strip().startswith("来源："):
                        i += 1
                    i += 1
                    continue
                else:
                    cap = doc.add_paragraph(style="表题")
                    render_inline(cap, cap_txt, size=10.5)
                    add_md_table(doc, rc3_table_rows(f"表{num}"))
                    stats["tables"] += 1
                    stats["captions"] += 1
                    if i + 1 < len(body_blocks) and body_blocks[i + 1]["type"] == "para" and \
                            body_blocks[i + 1]["text"].strip().startswith("来源："):
                        i += 1
                    i += 1
                    continue
            m2 = re.match(r"^\*\*(表\d+-\d+)(.*?)\*\*$", raw)
            if m2 and i + 1 < len(body_blocks) and body_blocks[i + 1]["type"] == "table":
                cap_txt = f"{m2.group(1)} {m2.group(2).strip()}".strip()
                mm = re.match(r"^表(\d)-(\d+)", cap_txt)
                if mm:
                    tbl_counters[mm.group(1)] = max(tbl_counters.get(mm.group(1), 0), int(mm.group(2)))
                p = doc.add_paragraph(style="表题")
                render_inline(p, cap_txt, size=10.5)
                add_md_table(doc, body_blocks[i + 1]["rows"])
                stats["tables"] += 1
                stats["captions"] += 1
                i += 2
                continue
            m3 = re.match(r"^\*\*(算法\d+-\d+)(.*?)\*\*$", raw)
            if m3 and i + 1 < len(body_blocks) and body_blocks[i + 1]["type"] == "code":
                cap_txt = f"{m3.group(1)} {m3.group(2).strip()}".strip()
                p = doc.add_paragraph(style="算法题注")
                render_inline(p, cap_txt, size=10.5)
                emit_code(doc, body_blocks[i + 1], stats)
                stats["algorithms"] += 1
                i += 2
                continue
            p = doc.add_paragraph()
            para_props(p, first_chars=2)
            render_inline(p, raw, size=12)
            stats["paras"] += 1
            i += 1
            continue
        if bt == "table":
            # uncaptioned table: assign sequential caption
            ch_num = chapter_num(current_chapter)
            tbl_counters[ch_num] = tbl_counters.get(ch_num, 0) + 1
            header0 = b["rows"][0][0] if b["rows"] and b["rows"][0] else ""
            title = UNCAPTIONED_TITLES.get((ch_num, header0), "数据表")
            cap_txt = f"表{ch_num}-{tbl_counters[ch_num]} {title}"
            p = doc.add_paragraph(style="表题")
            render_inline(p, cap_txt, size=10.5)
            add_md_table(doc, b["rows"])
            stats["tables"] += 1
            stats["captions"] += 1
            i += 1
            continue
        if bt == "code":
            emit_code(doc, b, stats)
            i += 1
            continue
        if bt == "eq":
            ch = re.sub(r"第", "", current_chapter)
            ch = ch[0] if ch else "4"
            ch_num = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7"}.get(ch, "4")
            eq_counters[ch_num] = eq_counters.get(ch_num, 0) + 1
            p = doc.add_paragraph(style="公式")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
            pf.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
            r = p.add_run("\t")
            set_run(r, size=12)
            add_omml(p, b["latex"])
            r = p.add_run(f"\t({ch_num}-{eq_counters[ch_num]})")
            set_run(r, size=12)
            stats["equations"] += 1
            i += 1
            continue
        if bt == "para":
            s = b["text"].strip()
            if in_refs and re.match(r"^\[\d+\]\s", s):
                p = doc.add_paragraph(style="参考文献")
                render_inline(p, reformat_ref(s), size=10.5)
                stats["refs"] += 1
                i += 1
                continue
            if re.match(r"^\d+\.\s", s):
                p = doc.add_paragraph()
                para_props(p, hanging_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            else:
                p = doc.add_paragraph()
                para_props(p, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            # 引理/定理/证明 label in 黑体
            m = re.match(r"^(\*\*?引理|引理|\*\*?定理|定理|\*\*?证明|证明)([^\n。]*[。]?)(\*\*)?(.*)$", s)
            if m and m.group(1).strip("*") in ("引理", "定理", "证明"):
                r = p.add_run(m.group(1).strip("*") + m.group(2))
                set_run(r, bold=True, east="黑体", size=12)
                render_inline(p, m.group(4), size=12)
            else:
                render_inline(p, s, size=12)
            # replace QED □ with ■
            for r in p.runs:
                if r.text and "□" in r.text:
                    r.text = r.text.replace("□", "■")
            stats["paras"] += 1
            i += 1
            continue
        i += 1

    if not chengguo_done:
        add_chengguo(doc, stats)
    set_update_fields(doc)
    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(json.dumps({"stats": stats}, ensure_ascii=False, indent=2))


def chapter_num(current: str) -> str:
    m = re.search(r"第([一二三四五六七1-7])章", current)
    if not m:
        return "4"
    c = m.group(1)
    return {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7"}.get(c, c)


def emit_code(doc, b: dict, stats: dict) -> None:
    p = doc.add_paragraph()
    para_props(p, line_pt=12, space_after=6)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "808080")
        pbdr.append(el)
    ppr.append(pbdr)
    for line in b["text"].splitlines():
        r = p.add_run(line)
        set_run(r, code=True, size=9)
        r.add_break()
        stats["manualBreaks"] += 1
    stats["paras"] += 1


def add_xiezhi(doc, stats) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = True
    p.add_run("致\u3000谢")
    body = doc.add_paragraph()
    para_props(body, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    render_inline(body,
                  "本论文的研究工作是在校内指导教师高建彬老师和校外指导教师王鹏老师的悉心指导下完成的。"
                  "两位老师在研究选题、方案设计与论文撰写等方面给予了大量帮助，在此致以诚挚的感谢。"
                  "感谢实验室同学在系统实现与实验过程中提供的支持，感谢家人在学业期间的理解与付出。",
                  size=12)
    stats["headings"] += 1
    stats["paras"] += 1


def add_chengguo(doc, stats) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = True
    p.add_run("攻读硕士学位期间取得的成果")
    body = doc.add_paragraph()
    para_props(body, first_chars=2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    render_inline(body, "[待填写]（如有与学位论文相关的发表论文、专利或获奖，按学校格式补充；如无，可删除本节。）", size=12)
    stats["headings"] += 1
    stats["paras"] += 1


if __name__ == "__main__":
    main()
