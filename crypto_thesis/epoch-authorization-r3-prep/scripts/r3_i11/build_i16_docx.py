# -*- coding: utf-8 -*-
"""I16: assemble the final thesis format-candidate DOCX from the master draft.

Pipeline:
  1. Apply I15 MINOR-1/MINOR-2 wording corrections to the master draft.
  2. Global citation renumbering by first appearance (GB/T 7714 sequential).
  3. Parse the corrected master into blocks (headings/paragraphs/tables/
     figures/equations/algorithms/references).
  4. Build a Word DOCX with real styles, native OMML equations, real tables,
     embedded frozen figures, auto TOC field, page numbering.
  5. Emit assembly-time audits (numeric/claim/citation/figure/table/equation/
     cross-reference) consumed by the I16 package generator.
"""
from __future__ import annotations

import copy
import io
import json
import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
MASTER = ROOT / "docs/thesis-integration/THESIS-INTEGRATED-MASTER-DRAFT-V1.md"
OUT_DIR = ROOT / "docs/final-manuscript"
OUTPUT = OUT_DIR / "output"
SRC_MD = OUT_DIR / "MASTER-SOURCE.md"
MML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
RC2_FIG = Path(r"D:\Research\crypto_thesis\epoch-authorization\docs\thesis-drafts\research-content-2-final\figures")
RC3_FIG = ROOT / "experiments/r3/formal/figures/i12-final"
RC3_TAB = ROOT / "experiments/r3/formal/tables/i12-final"
I12_RQ = ROOT / "docs/research-content-3-implementation/i12/formal-rq-results.json"

ACCESS_DATE = "2026-08-02"

EAST = "宋体"
HEI = "黑体"
LATIN = "Times New Roman"


# ---------------------------------------------------------------------------
# 1. I15 MINOR closures (wording only; no technical conclusion changes)
# ---------------------------------------------------------------------------

MINOR1_OLD = "已有研究已使用dyadic base表示任意区间并定义最小dyadic cover[2]；"
MINOR1_NEW = "已有研究关注分布式环境下 XML 等半结构化数据的分区与索引处理[2]；"

MINOR2_OLD = ("实验工件按照环境、配置、原始数据和处理脚本分层保存；"
              "这种组织与ACM关于文档化、完整且可执行工件的原则一致[5]，"
              "但本项目未申请或获得ACM工件徽章。")
MINOR2_NEW = ("实验工件按照环境、配置、原始数据和处理脚本分层保存，"
              "便于文档化、完整且可执行地复现；本项目未申请或获得ACM工件徽章。")


def apply_minor_fixes(text: str) -> str:
    fixes = [(MINOR1_OLD, MINOR1_NEW), (MINOR2_OLD, MINOR2_NEW)]
    for old, new in fixes:
        if old not in text:
            raise RuntimeError("I16: MINOR anchor not found: " + old[:40])
        text = text.replace(old, new)
    return text


# ---------------------------------------------------------------------------
# 2. Citation renumbering by first appearance
# ---------------------------------------------------------------------------


def compute_citation_map(text: str) -> tuple[str, dict[str, int]]:
    head, sep, tail = text.partition("## 参考文献")
    order: list[int] = []
    seen: set[int] = set()
    for m in re.finditer(r"\[(\d+)\]", head):
        k = int(m.group(1))
        if k not in seen:
            seen.add(k)
            order.append(k)
    mapping = {old: new for new, old in enumerate(order, 1)}

    def sub(m: re.Match) -> str:
        return f"[{mapping[int(m.group(1))]}]"

    head2 = re.sub(r"\[(\d+)\]", sub, head)
    # rebuild the reference section in new order, preserving trailing content
    tail_lines = tail.splitlines()
    ref_positions = [i for i, ln in enumerate(tail_lines) if re.match(r"^\[\d+\]\s", ln)]
    if ref_positions:
        last_ref = ref_positions[-1]
        entries = {}
        for i in ref_positions:
            m = re.match(r"^\[(\d+)\]\s+(.+)$", tail_lines[i])
            entries[int(m.group(1))] = m.group(2).rstrip()
        ordered = sorted(entries.items(), key=lambda kv: mapping.get(kv[0], kv[0]))
        refs = [f"[{mapping.get(old, old)}] {body}" for old, body in ordered]
        trailing = tail_lines[last_ref + 1:]
        tail2 = "## 参考文献\n\n" + "\n\n".join(refs) + "\n\n" + "\n".join(trailing).rstrip() + "\n"
    else:
        tail2 = tail
    return head2 + tail2, mapping


# ---------------------------------------------------------------------------
# 3. Block parser
# ---------------------------------------------------------------------------


def parse_blocks(text: str) -> list[dict]:
    lines = text.splitlines()
    blocks: list[dict] = []
    i = 0
    n = len(lines)
    while i < n:
        ln = lines[i].rstrip()
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        if s.startswith(">") or s.startswith("[LITERATURE_") or s.startswith("[文献覆盖说明") or s.startswith("[文献扩展建议"):
            i += 1
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", s)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2).strip()})
            i += 1
            continue
        m = re.match(r"^!\[(.*?)\]\((.*?)\)$", s)
        if m:
            blocks.append({"type": "image", "alt": m.group(1), "path": m.group(2)})
            i += 1
            continue
        if s.startswith("```"):
            lang = s[3:].strip()
            buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # closing fence
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(buf)})
            continue
        if s.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            if len(rows) >= 2:
                # drop separator row (|---|)
                data = [r for r in rows if not all(re.fullmatch(r":?-{3,}:?", c) for c in r)]
            else:
                data = rows
            blocks.append({"type": "table", "rows": data})
            continue
        if s.startswith("\\["):
            buf = [ln]
            i += 1
            while i < n and "\\]" not in lines[i]:
                buf.append(lines[i])
                i += 1
            if i < n:
                buf.append(lines[i])
                i += 1
            content = "\n".join(buf).strip()
            content = re.sub(r"^\\\[\s*", "", content)
            content = re.sub(r"\s*\\\]$", "", content)
            blocks.append({"type": "eq", "latex": content})
            continue
        if re.match(r"^\*\*(图|表|算法)", s):
            blocks.append({"type": "caption", "text": ln})
            i += 1
            continue
        blocks.append({"type": "para", "text": ln})
        i += 1
    return blocks


# ---------------------------------------------------------------------------
# 4. OMML (native Word equations)
# ---------------------------------------------------------------------------

_XSLT = None


def _transformer():
    global _XSLT
    if _XSLT is None:
        _XSLT = etree.XSLT(etree.parse(MML_XSL))
    return _XSLT


def latex_to_omml(latex: str) -> etree._Element:
    mathml = latex_to_mathml(latex, display="block")
    tree = etree.fromstring(mathml.encode("utf-8"))
    res = _transformer()(tree)
    root = res.getroot()
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    om = root.find(".//m:oMath", ns)
    if om is None:
        om = root
    return copy.deepcopy(om)


def add_omml(p, latex: str) -> None:
    p._p.append(latex_to_omml(latex))


# ---------------------------------------------------------------------------
# 5. Inline rendering
# ---------------------------------------------------------------------------

TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\s][^*]*?\*|`[^`]+?`|\\\(.*?\\\)|\\\[.*?\\\]|\[\d+\])", re.S)


def set_run(run, bold=False, italic=False, code=False, sup=False, size=None, east=EAST, latin=LATIN):
    if code:
        latin = "Consolas"
        east = "宋体"
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    run.bold = bold
    run.italic = italic
    if sup:
        run.font.superscript = True
    if size:
        run.font.size = Pt(size)


def render_inline(p, text: str, size: float = 12) -> None:
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2])
            set_run(r, bold=True, size=size)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1])
            set_run(r, code=True, size=size - 1)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = p.add_run(tok[1:-1])
            set_run(r, italic=True, size=size)
        elif tok.startswith("\\(") and tok.endswith("\\)"):
            add_omml(p, tok[2:-2])
        elif tok.startswith("\\[") and tok.endswith("\\]"):
            add_omml(p, tok[2:-2])
        elif re.fullmatch(r"\[\d+\]", tok):
            r = p.add_run(tok)
            set_run(r, sup=True, size=size)
        else:
            r = p.add_run(tok)
            set_run(r, size=size)


# ---------------------------------------------------------------------------
# 6. Word styles
# ---------------------------------------------------------------------------


def set_style_fonts(style, latin=LATIN, east=EAST, size=None, bold=None):
    style.font.name = latin
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_fonts(normal, size=12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)

    heads = [
        ("Heading 1", HEI, 16, True, 24, 24),
        ("Heading 2", HEI, 14, True, 18, 18),
        ("Heading 3", HEI, 12, True, 12, 12),
        ("Heading 4", HEI, 12, True, 12, 12),
    ]
    for name, east, size, bold, before, after in heads:
        st = doc.styles[name]
        set_style_fonts(st, east=east, size=size, bold=bold)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.2
        st.paragraph_format.keep_with_next = True

    for name, east, size, bold in [
        ("Caption", EAST, 10.5, True),
        ("Table Caption", EAST, 10.5, True),
        ("Algorithm Caption", EAST, 10.5, True),
        ("Reference", EAST, 10.5, False),
    ]:
        if name not in doc.styles:
            doc.styles.add_style(name, 1)  # paragraph style
        st = doc.styles[name]
        set_style_fonts(st, east=east, size=size, bold=bold)


def para_props(p, align=None, first_indent_chars=0, left_indent_cm=0.0, hanging_chars=0,
               keep_next=False, line=1.5, space_after=0):
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if first_indent_chars:
        ind.set(qn("w:firstLineChars"), str(first_indent_chars * 100))
        ind.set(qn("w:firstLine"), str(int(first_indent_chars * 240)))
    if hanging_chars:
        ind.set(qn("w:hangingChars"), str(hanging_chars * 100))
        ind.set(qn("w:hanging"), str(int(hanging_chars * 240)))
    if left_indent_cm:
        ind.set(qn("w:left"), str(int(left_indent_cm * 567)))
    if keep_next:
        pf.keep_with_next = True
    return p


# ---------------------------------------------------------------------------
# 7. Tables
# ---------------------------------------------------------------------------


def three_line_borders(table, header_rows=1):
    tbl = table._tbl
    tblPr = tbl.tblPr
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", 12), ("bottom", 12), ("left", 0), ("right", 0), ("insideH", 0), ("insideV", 0)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single" if sz else "none")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)
    # header bottom border
    for cell in table.rows[header_rows - 1].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "000000")
        tcB.append(bottom)
        tcPr.append(tcB)


def add_table(doc, rows, caption=None, table_caption=False, header_rows=1):
    if caption:
        p = doc.add_paragraph(style="Table Caption")
        para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
        render_inline(p, caption, size=10.5)
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width_cm = 15.0
    col_w = width_cm / ncols
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.width = Cm(col_w)
            txt = row[ci] if ci < len(row) else ""
            p0 = cell.paragraphs[0]
            p0.paragraph_format.line_spacing = 1.0
            p0.paragraph_format.space_after = Pt(2)
            p0.paragraph_format.space_before = Pt(2)
            render_inline(p0, txt, size=10.5)
            if ri < header_rows:
                for r in p0.runs:
                    r.bold = True
    three_line_borders(table, header_rows=header_rows)
    return table


# ---------------------------------------------------------------------------
# 8. RC3 frozen tables from JSON
# ---------------------------------------------------------------------------


def table_6_1_rows() -> list[list[str]]:
    d = json.load(io.open(RC3_TAB / "table-run-flow-eligibility.json", encoding="utf-8"))
    rows = d["rows"]
    by_exp: dict[str, list] = {}
    for r in rows:
        by_exp.setdefault(r["experiment"], []).append(r)
    out = [["实验", "配置数", "重复数", "有效运行", "VALID_SUCCESS", "VALID_EXPECTED_FAIL_CLOSED"]]
    for exp in sorted(by_exp):
        rs = by_exp[exp]
        valid = sum(1 for r in rs if r.get("valid"))
        succ = sum(1 for r in rs if r.get("disposition") == "VALID_SUCCESS")
        failc = sum(1 for r in rs if r.get("disposition") == "VALID_EXPECTED_FAIL_CLOSED")
        configs = len({(r["config"], r["repeat"]) for r in rs})
        repeats = len({r["repeat"] for r in rs})
        out.append([exp, str(len({r["config"] for r in rs})), str(repeats), str(valid), str(succ), str(failc)])
    total = len(rows)
    out.append(["合计", "-", "-", str(total),
                str(sum(1 for r in rows if r.get("disposition") == "VALID_SUCCESS")),
                str(sum(1 for r in rows if r.get("disposition") == "VALID_EXPECTED_FAIL_CLOSED"))])
    return out


def table_6_2_rows() -> list[list[str]]:
    d = json.load(io.open(I12_RQ, encoding="utf-8"))
    cards = d["cards"]
    out = [["配置", "n", "中位数/ms", "IQR/ms", "均值/ms", "95% Bootstrap CI/ms"]]
    for rq in ("RQ-2", "RQ-3"):
        levels = cards[rq].get("levels", {})
        for label in sorted(levels):
            v = levels[label]
            ci = v.get("ci95") or [None, None]
            ci_txt = f"[{ci[0]:.1f}, {ci[1]:.1f}]" if ci[0] is not None else "-"
            out.append([label.replace("recipient=", "接收者").replace("affected=", "资源").replace("body=", "Body "),
                        str(v["n"]), f"{v['median']:.1f}", f"{v['iqr']:.1f}", f"{v['mean']:.1f}", ci_txt])
    return out


def table_6_3_rows() -> list[list[str]]:
    d = json.load(io.open(RC3_TAB / "table-matched-local-kubo-recovery.json", encoding="utf-8"))
    cells = d["cells"]
    faults = ["NONE", "CORRUPT_RESTORE", "CID_MISMATCH", "BOTH_MISSING"]
    out = [["故障", "对象来源", "有效(n)", "恢复判定", "修复动作", "时长中位数/ms"]]
    for f in faults:
        for src in ("LOCAL_ONLY", "KUBO_REPLICA"):
            c = cells.get(f, {}).get(src)
            if not c:
                continue
            rec = c.get("recoveryDispositions") or {}
            rec_txt = "/".join(f"{k}:{v}" for k, v in rec.items()) or "-"
            rep = c.get("repairActions") or {}
            rep_txt = "/".join(f"{k}:{v}" for k, v in rep.items()) or "-"
            med = c.get("durationMedianMs")
            out.append([f, src, str(c["n"]), rec_txt, rep_txt, f"{med:.1f}" if med is not None else "-"])
    return out


def table_6_4_rows() -> list[list[str]]:
    d = json.load(io.open(RC3_TAB / "table-release-decision-outcome.json", encoding="utf-8"))
    dec = d["decisions"]
    out = [["释放判定", "运行数"], ["ALLOWED_AFTER_CURRENT_HEADER_ONLY", str(dec.get("ALLOWED_AFTER_CURRENT_HEADER_ONLY", 0))],
           ["DENIED", str(dec.get("DENIED", 0))], ["错误材料释放", str(d.get("wrongMaterialRelease", 0))]]
    return out


def table_6_5_rows() -> list[list[str]]:
    d = json.load(io.open(RC3_TAB / "table-environment-fingerprint.json", encoding="utf-8"))
    f = d.get("fingerprint", d)
    rows = [["项目", "值"]]
    pairs = [
        ("主机/角色", f.get("host", "") + "/" + f.get("role", "")),
        ("CPU", f.get("cpuModel", "")),
        ("物理/逻辑核", f"{f.get('physicalCores')}/{f.get('logicalCores')}"),
        ("内存", f"{f.get('ramBytes', 0) / 2**30:.2f} GiB"),
        ("OS/内核", f"{f.get('os')} {f.get('kernel')}"),
        ("虚拟化/网络", f"{f.get('virtualization')}；{f.get('network')}"),
        ("Python", f.get("pythonVersion", "")),
        ("Besu", f.get("besuVersion", "")),
        ("PostgreSQL", f.get("postgresqlVersion", "")),
        ("Kubo", f.get("kuboVersion", "")),
    ]
    for k, v in pairs:
        rows.append([k, str(v)])
    return rows


RC3_TABLES = {
    "表6-1": table_6_1_rows,
    "表6-2": table_6_2_rows,
    "表6-3": table_6_3_rows,
    "表6-4": table_6_4_rows,
    "表6-5": table_6_5_rows,
}

RC3_FIGS = {
    "图6-1": "fig-rq2-header-only-duration.png",
    "图6-2": "fig-rq3-body-rotation-duration.png",
    "图6-3": "fig-rq5-recovery-local-kubo.png",
}


# ---------------------------------------------------------------------------
# 9. Fields / TOC / page numbers
# ---------------------------------------------------------------------------


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:noProof"))
    ppr.append(rpr)

    def fld(tag: str, attrs=None, text=None):
        el = OxmlElement(tag)
        if attrs:
            for k, v in attrs.items():
                el.set(qn(k), v)
        if text:
            el.text = text
        return el

    r1 = OxmlElement("w:r"); r1.append(fld("w:fldChar", {"w:fldCharType": "begin"}))
    r2 = OxmlElement("w:r"); instr = fld("w:instrText", {"xml:space": "preserve"}, ' TOC \\o "1-3" \\h \\z \\u ')
    r2.append(instr)
    r3 = OxmlElement("w:r"); r3.append(fld("w:fldChar", {"w:fldCharType": "separate"}))
    r4 = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "（目录将在打开文档或更新域后自动生成）"; r4.append(t)
    r5 = OxmlElement("w:r"); r5.append(fld("w:fldChar", {"w:fldCharType": "end"}))
    for r in (r1, r2, r3, r4, r5):
        p._p.append(r)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    uf = settings.find(qn("w:updateFields"))
    if uf is None:
        uf = OxmlElement("w:updateFields")
        settings.append(uf)
    uf.set(qn("w:val"), "true")


def add_page_number_footer(section, fmt="decimal", start=1):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = str(start)
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))


def set_page(doc: Document):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# 10. DOCX assembly
# ---------------------------------------------------------------------------


def resolve_image(path: str) -> Path:
    p = Path(path.replace("/", "\\"))
    cands = []
    if p.is_absolute():
        cands.append(p)
    else:
        cands.append(ROOT / p)
        cands.append(ROOT.parent / p)
        cands.append(RC2_FIG / p.name)
    for c in cands:
        if c.is_file():
            return c
    return p


def chapter_key(text: str) -> str | None:
    t = text.strip()
    m = re.match(r"^(第[一二三四五六七]章|参考文献|附录A)", t)
    return m.group(1) if m else None


def heading_level_of(text: str) -> int | None:
    t = text.strip()
    if chapter_key(t):
        return 1
    m = re.match(r"^(\d+)\.(\d+)(?:\.(\d+))?\s", t)
    if m:
        return 2 if not m.group(3) else 3
    return None


def normalize_heading(text: str) -> str:
    return text.replace("第5章", "第五章")


def build_docx(blocks: list[dict], out_path: Path) -> dict:
    doc = Document()
    setup_styles(doc)
    set_page(doc)

    # ---- cover ----
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("电子科技大学")
    set_run(r, size=26, east=HEI, latin="SimHei")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("专业学位硕士学位论文")
    set_run(r, size=22, east=HEI, latin="SimHei")
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《面向非连续时间约束的区块链数据共享关键技术研究及实现》")
    set_run(r, size=18, east=HEI, latin="SimHei")
    for _ in range(3):
        doc.add_paragraph()
    cover_fields = [
        ("学    号", "202422081113"),
        ("姓    名", "王  威"),
        ("学    院", "计算机科学与工程学院（网络空间安全学院）"),
        ("专业学位类别", "计算机技术"),
        ("校内指导教师", "高建彬"),
        ("校外指导教师", "王  鹏"),
        ("提交日期", "[待填写]"),
    ]
    for k, v in cover_fields:
        p = doc.add_paragraph()
        para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.8)
        r = p.add_run(f"{k}：{v}")
        set_run(r, size=14, east=EAST)
    doc.add_page_break()
    add_page_number_footer(doc.sections[0], fmt="lowerRoman", start=1)

    stats = {"paras": 0, "tables": 0, "figures": 0, "equations": 0, "refs": 0,
             "algorithms": 0, "headings": 0, "captions": 0}
    eq_counters: dict[str, int] = {}
    current_chapter = "前置"
    front_done = False
    skip_toc_summary = False

    i = 0
    while i < len(blocks):
        b = blocks[i]
        bt = b["type"]

        if bt == "heading":
            text = b["text"]
            if text.startswith("《") or "集成母本候选稿" in text:
                i += 1
                continue
            key = chapter_key(text)
            if text == "中文摘要" or text == "Abstract":
                p = doc.add_paragraph(style="Heading 1")
                para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                if text == "Abstract":
                    p.paragraph_format.page_break_before = True
                p.add_run(text)
                stats["headings"] += 1
                i += 1
                continue
            if text == "关键词" or text == "Keywords":
                p = doc.add_paragraph(style="Heading 1")
                para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                p.add_run(text)
                stats["headings"] += 1
                i += 1
                continue
            if text == "目录":
                p = doc.add_paragraph(style="Heading 1")
                para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                p.paragraph_format.page_break_before = True
                p.add_run("目录")
                stats["headings"] += 1
                add_toc(doc)
                skip_toc_summary = True
                i += 1
                continue
            if key:
                text = normalize_heading(text)
                if not front_done:
                    front_done = True
                    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                    add_page_number_footer(new_section, fmt="decimal", start=1)
                    set_page(doc)
                p = doc.add_paragraph(style="Heading 1")
                para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                p.paragraph_format.page_break_before = True
                p.add_run(text)
                current_chapter = text
                stats["headings"] += 1
                i += 1
                continue
            lvl = heading_level_of(text)
            if lvl is None:
                lvl = min(b["level"], 4)
            p = doc.add_paragraph(style=f"Heading {lvl}")
            para_props(p, keep_next=True, line=1.2)
            p.add_run(text)
            stats["headings"] += 1
            i += 1
            continue

        if bt == "image":
            path = resolve_image(b["path"])
            alt = b["alt"].strip()
            p = doc.add_paragraph()
            para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.0)
            run = p.add_run()
            try:
                run.add_picture(str(path), width=Cm(14.5))
            except Exception:
                run.add_text(f"[图片缺失: {b['path']}]")
            cap = doc.add_paragraph(style="Caption")
            para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
            render_inline(cap, alt, size=10.5)
            stats["figures"] += 1
            stats["captions"] += 1
            # skip duplicate caption paragraph immediately following
            if i + 1 < len(blocks) and blocks[i + 1]["type"] == "para":
                nxt = blocks[i + 1]["text"].strip()
                if nxt in (alt, alt + "。"):
                    i += 1
            i += 1
            continue

        if bt == "caption":
            raw = b["text"].strip()
            m_rc3 = re.match(r"^\*\*(图|表)\s*(\d+-\d+)([^*]*?)\*\*(.*)$", raw)
            if m_rc3 and m_rc3.group(2).startswith("6-"):
                kind, num, inner, rest = m_rc3.groups()
                cap_txt = f"{kind}{num}{inner} {rest}".strip()
                if kind == "图":
                    fname = RC3_FIGS.get(f"图{num}")
                    if fname:
                        p = doc.add_paragraph()
                        para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.0)
                        run = p.add_run()
                        run.add_picture(str(RC3_FIG / fname), width=Cm(14.5))
                        stats["figures"] += 1
                    cap = doc.add_paragraph(style="Caption")
                    para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                    clean = re.sub(r"\s+", " ", cap_txt.replace("（E2）", "").replace("（E3）", "").replace("（E5）", "")).strip()
                    render_inline(cap, clean, size=10.5)
                    stats["captions"] += 1
                    if i + 1 < len(blocks) and blocks[i + 1]["type"] == "para" and blocks[i + 1]["text"].strip().startswith("来源："):
                        i += 1
                    i += 1
                    continue
                else:
                    cap = doc.add_paragraph(style="Table Caption")
                    para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                    render_inline(cap, cap_txt, size=10.5)
                    rows = RC3_TABLES[f"表{num}"]()
                    add_table(doc, rows)
                    stats["tables"] += 1
                    stats["captions"] += 1
                    if i + 1 < len(blocks) and blocks[i + 1]["type"] == "para" and blocks[i + 1]["text"].strip().startswith("来源："):
                        i += 1
                    i += 1
                    continue
            # generic caption: table 4-x/5-x or algorithm caption
            m2 = re.match(r"^\*\*(表\d+-\d+)(.*?)\*\*$", raw)
            if m2:
                cap_txt = f"{m2.group(1)} {m2.group(2).strip()}".strip()
                if i + 1 < len(blocks) and blocks[i + 1]["type"] == "table":
                    cap = doc.add_paragraph(style="Table Caption")
                    para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                    render_inline(cap, cap_txt, size=10.5)
                    add_table(doc, blocks[i + 1]["rows"])
                    stats["tables"] += 1
                    stats["captions"] += 1
                    i += 2
                    continue
            m3 = re.match(r"^\*\*(算法\d+-\d+)(.*?)\*\*$", b["text"].strip())
            if m3:
                cap_txt = f"{m3.group(1)} {m3.group(2).strip()}".strip()
                if i + 1 < len(blocks) and blocks[i + 1]["type"] == "code":
                    cap = doc.add_paragraph(style="Algorithm Caption")
                    para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, keep_next=True, line=1.2)
                    render_inline(cap, cap_txt, size=10.5)
                    emit_code(doc, blocks[i + 1], stats)
                    stats["algorithms"] += 1
                    i += 2
                    continue
            # fallback: treat as normal paragraph
            p = doc.add_paragraph()
            para_props(p, first_indent_chars=2)
            render_inline(p, b["text"], size=12)
            stats["paras"] += 1
            i += 1
            continue

        if bt == "table":
            add_table(doc, b["rows"])
            stats["tables"] += 1
            i += 1
            continue

        if bt == "code":
            emit_code(doc, b, stats)
            i += 1
            continue

        if bt == "eq":
            ch = re.sub(r"第", "", current_chapter)
            ch = ch[0] if ch else "4"
            try:
                ch_num = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5", "六": "6", "七": "7"}.get(ch, "4")
            except Exception:
                ch_num = "4"
            eq_counters[ch_num] = eq_counters.get(ch_num, 0) + 1
            p = doc.add_paragraph()
            para_props(p, line=1.4)
            pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Cm(7.75), WD_TAB_ALIGNMENT.CENTER)
            pf.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
            r = p.add_run("\t")
            set_run(r, size=12)
            add_omml(p, b["latex"])
            r = p.add_run(f"\t({ch_num}-{eq_counters[ch_num]})")
            set_run(r, size=12)
            stats["equations"] += 1
            i += 1
            continue

        if bt == "para":
            text = b["text"].strip()
            if skip_toc_summary:
                skip_toc_summary = False
                i += 1
                continue
            if text.startswith("来源：") or text == "Keywords" or text == "关键词":
                i += 1
                continue
            if current_chapter == "参考文献" and re.match(r"^\[\d+\]\s", text):
                p = doc.add_paragraph()
                para_props(p, hanging_chars=2, line=1.5)
                mnum = re.match(r"^(\[\d+\])(.*)$", text)
                r = p.add_run(mnum.group(1))
                set_run(r, size=10.5)
                render_inline(p, mnum.group(2), size=10.5)
                stats["refs"] += 1
                i += 1
                continue
            if re.match(r"^\d+\.\s", text):
                p = doc.add_paragraph()
                para_props(p, left_indent_cm=0.74, line=1.5)
                render_inline(p, text, size=12)
            else:
                p = doc.add_paragraph()
                para_props(p, first_indent_chars=2)
                render_inline(p, text, size=12)
            stats["paras"] += 1
            i += 1
            continue

        i += 1

    set_update_fields(doc)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return stats


def emit_code(doc, b: dict, stats: dict) -> None:
    lang = b.get("lang", "")
    text = b["text"].rstrip()
    if lang == "mermaid":
        # text diagram (图5-A / 图5-B)
        p = doc.add_paragraph()
        para_props(p, line=1.0)
        ppr = p._p.get_or_add_pPr()
        ppr.append(box_borders())
        for line in text.splitlines():
            r = p.add_run(line)
            set_run(r, code=True, size=9)
            r.add_break()
        stats["paras"] += 1
        return
    p = doc.add_paragraph()
    para_props(p, line=1.0)
    ppr = p._p.get_or_add_pPr()
    ppr.append(box_borders())
    for line in text.splitlines():
        r = p.add_run(line)
        set_run(r, code=True, size=9)
        r.add_break()
    stats["paras"] += 1


def box_borders() -> OxmlElement:
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "808080")
        pbdr.append(el)
    return pbdr


# ---------------------------------------------------------------------------
# 11. Audits
# ---------------------------------------------------------------------------


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def audit_numeric(master_text: str, docx_text: str) -> dict:
    keys = ["15120", "145", "35", "29", "120", "25", "2430", "9720", "77760", "233280",
            "103680", "324", "108", "64 KiB", "8 MiB", "17.78", "17.93", "5115", "5083", "6696"]
    missing = [k for k in keys if k in master_text and k not in docx_text]
    return {"keysChecked": len(keys), "missingInDocx": missing, "drift": len(missing)}


def audit_claims(text: str) -> dict:
    patterns = ["首次", "首个", "国际领先", "填补空白", "显著提升", "领先水平", "最优", "完美解决",
                "完全杜绝", "绝对安全", "优于所有", "现有研究缺乏", "尚未解决"]
    negations = ["不存在", "不等同于", "不预设", "不得", "不声称", "不能", "没有", "不保证",
                 "不构成", "不等于", "无法", "难以", "不能证明", "未", "不是", "并不"]
    hits: dict[str, list[str]] = {}
    for k in patterns:
        found = []
        for m in re.finditer(re.escape(k), text):
            start = max(0, m.start() - 25)
            ctx = text[start:m.end() + 10]
            if any(neg in ctx for neg in negations):
                continue  # negated/quoted boundary, not an overclaim
            found.append(ctx[:60].replace("\n", " "))
        if found:
            hits[k] = found
    return {"forbiddenHits": hits, "count": sum(len(v) for v in hits.values())}


def audit_citations(docx_text: str) -> dict:
    head, _, tail = docx_text.partition("参考文献")
    cited = sorted({int(m) for m in re.findall(r"\[(\d+)\]", head)})
    refs = len(re.findall(r"(?m)^\[\d+\] ", tail))
    missing = [k for k in cited if k > refs]
    return {"citedKeys": cited, "referenceEntries": refs, "missing": missing,
            "orphan": [k for k in range(1, refs + 1) if k not in cited]}


def audit_figures_tables(doc) -> dict:
    figs = len(doc.inline_shapes)
    tables = len(doc.tables)
    text = "\n".join(p.text for p in doc.paragraphs)
    fig_caps = sorted(set(re.findall(r"图(\d+-\d+)", text)))
    tbl_caps = sorted(set(re.findall(r"表(\d+-\d+)", text)))
    return {"inlineImages": figs, "tables": tables, "figureCaptions": fig_caps, "tableCaptions": tbl_caps}


def audit_equations(doc) -> dict:
    xml = etree.tostring(doc.element.body, encoding="unicode")
    return {"ommlCount": xml.count("<m:oMath"), "equationNumbers": len(re.findall(r"\([4-6]-\d+\)", xml))}


def audit_crossrefs(docx_text: str) -> dict:
    caps_fig = set(re.findall(r"图(\d+-\d+)", docx_text))
    caps_tbl = set(re.findall(r"表(\d+-\d+)", docx_text))
    refs_fig = set(re.findall(r"图\s*(\d+-\d+)", docx_text))
    refs_tbl = set(re.findall(r"表\s*(\d+-\d+)", docx_text))
    return {"figureKeys": sorted(caps_fig), "tableKeys": sorted(caps_tbl),
            "figureRefs": sorted(refs_fig), "tableRefs": sorted(refs_tbl)}


def main() -> None:
    created = datetime.now(timezone.utc).isoformat()
    text = io.open(MASTER, encoding="utf-8").read()
    text = apply_minor_fixes(text)
    text2, citation_map = compute_citation_map(text)
    SRC_MD.parent.mkdir(parents=True, exist_ok=True)
    io.open(SRC_MD, "w", encoding="utf-8").write(text2)

    blocks = parse_blocks(text2)
    out_path = OUTPUT / "THESIS-FORMAT-CANDIDATE-V1.docx"
    stats = build_docx(blocks, out_path)

    doc = Document(str(out_path))
    docx_text = extract_docx_text(out_path)
    audits = {
        "numeric": audit_numeric(text2, docx_text),
        "claims": audit_claims(docx_text),
        "citations": audit_citations(docx_text),
        "figuresTables": audit_figures_tables(doc),
        "equations": audit_equations(doc),
        "crossrefs": audit_crossrefs(docx_text),
    }
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    io.open(OUT_DIR / "assembly-audits.json", "w", encoding="utf-8").write(
        json.dumps({"schemaVersion": "I16AssemblyAuditsV1", "generatedAt": created,
                    "citationMap": citation_map, "stats": stats, "audits": audits},
                   ensure_ascii=False, indent=2))
    print(json.dumps({"stats": stats, "audits": audits}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
