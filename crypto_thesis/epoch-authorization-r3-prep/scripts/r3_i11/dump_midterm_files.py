"""Dump structure of the official midterm template and the reference report."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


TEMPLATE = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\附件2：专业学位研究生学位论文中期考评表-2023版.docx")
REFERENCE = Path(r"D:\Users\wangw\Desktop\中期和小论文\shy-专业学位研究生学位论文中期考评表.docx")


def dump(path: Path, label: str) -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print(f"\n################ {label} ################")
    doc = Document(str(path))
    print("sections:", len(doc.sections), "| tables:", len(doc.tables))
    print("---- paragraphs ----")
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            style = p.style.name if p.style else "?"
            print(f"[{i}] ({style}) {t[:120]}")
    print("---- tables ----")
    for ti, t in enumerate(doc.tables):
        print(f"table {ti}: rows={len(t.rows)} cols={len(t.columns)}")
        for ri, row in enumerate(t.rows[:20]):
            cells = " || ".join(c.text.strip().replace("\n", " ")[:34] for c in row.cells)
            print(f"  r{ri}: {cells[:200]}")
        if len(t.rows) > 20:
            print(f"  ... ({len(t.rows)} rows total)")


def main() -> None:
    dump(TEMPLATE, "OFFICIAL EMPTY TEMPLATE")
    dump(REFERENCE, "REFERENCE REPORT (WRITING_REFERENCE)")


if __name__ == "__main__":
    main()
