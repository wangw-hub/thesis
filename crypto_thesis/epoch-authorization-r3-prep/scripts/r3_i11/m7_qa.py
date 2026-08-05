# -*- coding: utf-8 -*-
"""M7 structural QA against the built DOCX and M7 source."""
from __future__ import annotations

import json
import re
import sys

from docx import Document
from lxml import etree


ROOT = r"D:\Research\crypto_thesis\epoch-authorization-r3-prep"
DOCX = ROOT + r"\docs\midterm-report\m7\output\王威-专业学位研究生学位论文中期考评表-M7最终候选稿.docx"
SRC = ROOT + r"\docs\midterm-report\m7\M7-MIDTERM-SOURCE.md"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    report: dict = {}

    # ---------------- source-level checks ----------------
    src = open(SRC, encoding="utf-8").read()
    src_body = src[: src.find("### 参考文献")]
    tokens = re.findall(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]", src_body)
    first_seen: dict[int, int] = {}
    order_err = []
    last_pos = -1
    for m in re.finditer(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]", src_body):
        nums = []
        for part in m.group(0)[1:-1].replace(" ", "").split(","):
            if "-" in part:
                a, b = part.split("-")
                nums.extend(range(int(a), int(b) + 1))
            else:
                nums.append(int(part))
        for n in nums:
            if n not in first_seen:
                first_seen[n] = m.start()
    seq = sorted(first_seen.items(), key=lambda kv: kv[1])
    nums = [n for n, _ in seq]
    if nums != list(range(1, len(nums) + 1)):
        order_err.append(f"citation numbers not contiguous: {nums[:60]}")
    report["source_citation_count"] = len(nums)
    report["citation_order_error"] = order_err

    ref_block = src[src.find("### 参考文献"): src.find("### 4．阶段性研究成果")]
    refs = re.findall(r"^\[(\d+)\]", ref_block, re.M)
    report["reference_count"] = len(refs)
    report["reference_numbers"] = [int(x) for x in refs]
    years = re.findall(r", (\d{4})(?:,|\[|$|\))", ref_block)
    report["reference_years_sample"] = years
    years_int = []
    for m in re.finditer(r"(?:J\]|C\]|S\]|EB/OL\]|M\]).*?(\d{4})", ref_block, re.S):
        years_int.append(int(m.group(1)))
    # fallback: years near the end of each reference entry
    entries = re.split(r"\n\n\[", ref_block)
    y2021 = y2024 = 0
    for e in entries:
        ym = re.findall(r"(20\d\d)", e)
        if not ym:
            continue
        y = int(ym[-1])
        if y >= 2021:
            y2021 += 1
        if y >= 2024:
            y2024 += 1
    report["ref_2021_2026_count"] = y2021
    report["ref_2024_2026_count"] = y2024
    report["ref_2021_2026_ratio"] = round(y2021 / len(refs), 3) if refs else 0

    # orphan check: every [n] in the reference list cited in body
    cited = []
    for m in re.finditer(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]", src_body):
        for part in m.group(0)[1:-1].replace(" ", "").split(","):
            if "-" in part:
                a, b = part.split("-")
                cited.extend(range(int(a), int(b) + 1))
            else:
                cited.append(int(part))
    cited_set = set(cited)
    orphans = [n for n in range(1, len(refs) + 1) if n not in cited_set]
    report["orphan_references"] = orphans

    # algorithm checks from source
    algos = re.findall(r"\[算法块：算法(\d+) ([^\n]+)", src)
    report["algorithms"] = [f"{a} {b}" for a, b in algos]
    report["algo_numbering_ok"] = [int(a) for a, _ in algos] == list(range(1, 9))
    idx6 = next(i for i, (a, _) in enumerate(algos) if a == "6")
    idx7 = next(i for i, (a, _) in enumerate(algos) if a == "7")
    report["algo6_before_algo7"] = idx6 < idx7
    report["formula_count_source"] = len(re.findall(r"^\[公式：", src, re.M))

    # ---------------- DOCX-level checks ----------------
    doc = Document(DOCX)
    body = doc.element.body

    # formula placeholders
    empty_nary = 0
    formula_paras = 0
    for om in body.iter(M + "oMath"):
        xml = etree.tostring(om, encoding="unicode")
        if "<m:e/>" in xml:
            empty_nary += 1
        formula_paras += 1
    report["docx_oMath_count"] = formula_paras
    report["docx_empty_nary"] = empty_nary

    # algorithm end marker "算法结束" must not be visible in DOCX text
    all_text = "\n".join(p.text or "" for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += "\n" + (cell.text or "")
    report["algo_end_text_in_docx"] = "算法结束" in all_text

    # tables: top-level and nested; tblHeader/cantSplit on every data table
    tables = doc.tables
    nested = []
    for tbl in tables:
        nested.extend(tbl._tbl.findall(f".//{W}tbl"))
    all_tbls = [t._tbl for t in tables] + nested
    report["docx_tables"] = len(tables) + len(nested)
    header_flags = []
    for ti, tbl in enumerate(all_tbls):
        hdr = tbl.findall(f".//{W}trPr/{W}tblHeader")
        cant = tbl.findall(f".//{W}trPr/{W}cantSplit")
        header_flags.append({"table": ti, "tblHeader": len(hdr), "cantSplit": len(cant)})
    report["table_header_flags"] = header_flags

    # cover double-check: degree-level paragraph should have exactly one ☑
    check_chars = 0
    for para in doc.paragraphs:
        if "攻读学位级别" in (para.text or ""):
            check_chars = (para.text or "").count("☑")
    report["cover_checked_marks"] = check_chars
    # Wingdings syms remaining on degree line
    sym_count = 0
    for para in doc.paragraphs:
        if "攻读学位级别" in (para.text or ""):
            sym_count = len(para._p.findall(f".//{W}sym"))
    report["cover_wingdings_sym"] = sym_count

    # figures count
    figs = len(body.findall(f".//{W}drawing")) + len(body.findall(f".//{W}pict"))
    report["docx_figures"] = figs

    print(json.dumps(report, ensure_ascii=False, indent=1))


if __name__ == "__main__":
    main()
