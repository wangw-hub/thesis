# -*- coding: utf-8 -*-
"""M7: dump OMML for equations (1) and (3), check placeholders; probe template pagination."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
M6_DOCX = ROOT / "docs/midterm-report/m6/output/王威-专业学位研究生学位论文中期考评表-M6候选稿.docx"
BLANK = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\附件2：专业学位研究生学位论文中期考评表-2023版.docx")
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(M6_DOCX))
    cell = doc.tables[0].rows[4].cells[0]
    count = 0
    for p in cell.paragraphs:
        if re.match(r"^\(\d+\)", p.text.strip()):
            count += 1
            if count in (1, 3):
                print("=" * 20, "equation", count, "text:", repr(p.text[:50]))
                om = p._p.findall(f".//{M}oMath")
                for o in om[:1]:
                    xml = etree.tostring(o, encoding="unicode")
                    # strip namespaces for readability
                    xml = re.sub(r"\sxmlns:[a-zA-Z0-9]+=\"[^\"]*\"", "", xml)
                    print(xml[:2600])
                # check placeholder
                full = etree.tostring(p._p, encoding="unicode")
                print("placeholder tokens:", [t for t in ("w:placeholder", "\ufffd", "MISSING") if t in full])
    print("display equations:", count)

    # blank template pagination probe
    print("=" * 20, "blank template page breaks")
    bd = Document(str(BLANK))
    body = bd.element.body
    breaks = 0
    for br in body.iter(f"{W}br"):
        if br.get(f"{W}type") == "page":
            breaks += 1
            print("explicit page break found")
    sects = body.findall(f"{W}sectPr")
    print("section breaks:", len(sects))
    for s in sects:
        pg = s.find(f"{W}pgSz")
        print("pgSz:", pg.get(f"{W}w") if pg is not None else None, pg.get(f"{W}h") if pg is not None else None)
    # count paragraphs between cover and first table
    print("tables:", len(bd.tables))


if __name__ == "__main__":
    main()
