"""Full-cell dump of the cover/flyleaf template + manual-break audit of V1."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


COVER = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\电子科技大学研究生学位论文封面及扉页 - 适用于专业学位硕士_081705087525.docx")
V1 = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep\docs\final-manuscript\output\THESIS-FORMAT-CANDIDATE-V1.docx")


def cell_runs_info(cell) -> str:
    parts = []
    for p in cell.paragraphs:
        for r in p.runs:
            sz = r.font.size.pt if r.font.size else None
            name = r.font.name
            ea = None
            rpr = r._element.find(qn("w:rPr"))
            if rpr is not None:
                rf = rpr.find(qn("w:rFonts"))
                if rf is not None:
                    ea = rf.get(qn("w:eastAsia"))
            parts.append(f"{r.text!r}[sz={sz},f={name},ea={ea},b={r.bold}]")
    return " ".join(parts)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(COVER))
    for ti, t in enumerate(doc.tables):
        print(f"===== TABLE {ti} =====")
        seen = set()
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if not txt:
                    continue
                key = (txt, ci)
                if key in seen:
                    continue
                seen.add(key)
                print(f"  r{ri}c{ci}: {txt[:90]!r}")
                info = cell_runs_info(cell)
                if info:
                    print(f"        runs: {info[:200]}")

    print()
    print("===== V1 MANUAL BREAK AUDIT =====")
    v1 = Document(str(V1))
    body = v1.element.body
    xml = body.xml
    breaks = re.findall(r"<w:br\b[^>]*/>", xml)
    print("total <w:br/> in V1 document.xml:", len(breaks))
    from collections import Counter
    kinds = Counter(re.search(r'w:type="([^"]+)"', b).group(1) if re.search(r'w:type="([^"]+)"', b) else "text_wrap" for b in breaks)
    print("break kinds:", dict(kinds))
    # paragraphs containing breaks (body paragraphs vs table cells)
    para_breaks = 0
    table_breaks = 0
    for p in v1.paragraphs:
        if p._p.xml.count("<w:br"):
            para_breaks += 1
    for t in v1.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    if p._p.xml.count("<w:br"):
                        table_breaks += 1
    print("paragraphs with manual breaks (body, non-table):", para_breaks)
    print("table-cell paragraphs with manual breaks:", table_breaks)
    # one-sentence-per-paragraph scan on body text
    body_text = [p.text for p in v1.paragraphs]
    short = sum(1 for t in body_text if 0 < len(t) < 60 and not t.startswith(("图", "表", "算法", "第", "4.", "5.", "6.")))
    print("short body paragraphs (<60 chars, non-caption):", short)


if __name__ == "__main__":
    main()
