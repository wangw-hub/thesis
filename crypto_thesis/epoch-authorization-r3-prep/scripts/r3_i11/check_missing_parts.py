"""File-based check for appendix and mermaid captions in DOCX/PDF."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
DOCX = ROOT / "docs/final-manuscript/output/THESIS-FORMAT-CANDIDATE-V1.docx"
PDF = ROOT / "docs/final-manuscript/output/THESIS-FORMAT-CANDIDATE-V1.pdf"

KEYS = ["附录A", "复现说明", "图5-A", "图5-B", "stateDiagram", "sequenceDiagram",
        "accepted runs", "chainId 2026080201"]


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(DOCX))
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            docx_text += "\n" + " | ".join(c.text for c in row.cells)
    print("===== DOCX =====")
    for k in KEYS:
        print(f"{k}: {docx_text.count(k)}")
    # tail of docx body
    idx = docx_text.find("复现说明")
    print("DOCX appendix context:", docx_text[idx - 100:idx + 400] if idx >= 0 else "N/A")

    import pypdfium2 as pdfium
    d = pdfium.PdfDocument(str(PDF))
    pdf_text = "\n".join(d[i].get_textpage().get_text_range() for i in range(len(d)))
    print("===== PDF =====")
    for k in KEYS:
        print(f"{k}: {pdf_text.count(k)}")
    print("PDF tail:", pdf_text[-500:])


if __name__ == "__main__":
    main()
