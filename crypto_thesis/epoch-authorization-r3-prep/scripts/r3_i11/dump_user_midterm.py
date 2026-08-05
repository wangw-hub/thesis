"""Dump the user's own midterm form, progress report, and small paper."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


FORM = Path(r"D:\Users\wangw\Desktop\中期和小论文\王威专业学位研究生学位论文中期考评表.docx")
REPORT = Path(r"D:\Users\wangw\Desktop\中期和小论文\王威-专业学位研究生学位论文中期考核研究进展报告-完整定稿版.docx")
PAPER = Path(r"D:\Users\wangw\Desktop\中期和小论文\学位论文小论文.docx")


def dump(path: Path, label: str, max_paras: int = 200, max_cell: int = 3000) -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print(f"\n################ {label} ################")
    doc = Document(str(path))
    print("sections:", len(doc.sections), "| tables:", len(doc.tables))
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            print(f"[{i}] {t[:150]}")
        if i > max_paras:
            break
    for ti, t in enumerate(doc.tables):
        print(f"\n===== table {ti}: rows={len(t.rows)} =====")
        for ri, row in enumerate(t.rows):
            seen = set()
            texts = []
            for c in row.cells:
                if id(c._tc) in seen:
                    continue
                seen.add(id(c._tc))
                texts.append(c.text)
            for ci, txt in enumerate(texts):
                print(f"--- r{ri}c{ci} len={len(txt)} ---")
                print(txt[:max_cell])


def main() -> None:
    dump(FORM, "USER OWN MIDTERM FORM")
    dump(REPORT, "USER PROGRESS REPORT 完整定稿版")
    dump(PAPER, "SMALL PAPER (学位论文小论文)", max_paras=40, max_cell=800)


if __name__ == "__main__":
    main()
