# -*- coding: utf-8 -*-
import re
import sys
from pathlib import Path

from docx import Document


OUT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep\docs\midterm-report\m3\output\王威-专业学位研究生学位论文中期考评表-M3候选稿.docx")


def main():
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(OUT))
    xml = OUT.read_bytes().decode("utf-8", errors="ignore")
    print("doc tables:", len(doc.tables))
    print("images:", len(doc.inline_shapes))
    print("oMath count:", xml.count("<m:oMath"))
    print("supVertAlign:", xml.count('w:vertAlign w:val="superscript"'))
    print("firstLineChars=200:", xml.count('w:firstLineChars="200"'))
    print("hangingChars:", xml.count("w:hangingChars"))
    print("pBdr (algo boxes):", xml.count("<w:pBdr>"))
    texts = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                texts.append(cell.text)
    full = "\n".join(texts)
    print("total chars in tables:", len(full))
    for fig in ["图1 ", "图5 ", "图8 ", "图12 ", "图18 ", "图20 "]:
        print(f"has '{fig.strip()}':", fig.strip() in full)
    for ref in ["[1] Bertino", "[29] Gray"]:
        print(f"ref {ref}:", ref in full)
    for alg in ["算法1 非连续时间策略规范化", "算法8 RecoveryCoordinator"]:
        print(f"algo {alg}:", alg in full)
    for eqn in ["(1)", "(13)", "(26)"]:
        print(f"eqn {eqn}:", eqn in full)


if __name__ == "__main__":
    main()
