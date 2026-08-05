# -*- coding: utf-8 -*-
"""M6: inspect built DOCX structure."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
OUT = ROOT / "docs/midterm-report/m6/output/王威-专业学位研究生学位论文中期考评表-M6候选稿.docx"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(OUT))
    t0 = doc.tables[0]
    cell = t0.rows[4].cells[0]
    print("nested tables:", len(cell.tables))
    for i, tb in enumerate(cell.tables):
        print("table", i, "rows:", len(tb.rows), "cols:", len(tb.columns),
              "first row:", " | ".join(c.text[:14] for c in tb.rows[0].cells))
    # count equations
    om = cell._tc.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")
    print("oMath:", len(om))
    blips = cell._tc.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
    print("images:", len(blips))
    # check algorithm paragraphs
    paras = cell.paragraphs
    algos = [p.text[:50] for p in paras if p.text.strip().startswith("算法")]
    print("algo-ish paragraphs:", len(algos))
    for a in algos:
        print(" |", a)


if __name__ == "__main__":
    main()
