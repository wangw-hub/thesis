# -*- coding: utf-8 -*-
"""M6: inspect M5 DOCX rendering of algorithm blocks and equations."""
from __future__ import annotations

import io
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
OUT = ROOT / "docs/midterm-report/m5/output/王威-专业学位研究生学位论文中期考评表-M5候选稿.docx"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(OUT))
    t0 = doc.tables[0]
    cell = t0.rows[4].cells[0]
    paras = cell.paragraphs
    print("paragraphs in progress cell:", len(paras))
    algo_idx = []
    eq_idx = []
    for i, p in enumerate(paras):
        tx = p.text
        if tx.strip().startswith("算法"):
            algo_idx.append((i, tx[:60]))
        if re.match(r"^\(\d+\)", tx.strip()) and "oMath" in p._p.xml:
            eq_idx.append((i, tx.strip()[:40]))
    print("algorithm-like paragraphs:", len(algo_idx))
    for i, t in algo_idx[:10]:
        print(i, "|", t)
    print("equation-like paragraphs:", len(eq_idx))
    for i, t in eq_idx[:30]:
        print(i, "|", t)

    # count oMath elements
    om = cell._cell.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")
    print("oMath elements in progress cell:", len(om))
    # count inline pictures
    blips = cell._cell.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
    print("blips (images):", len(blips))
    # count tables
    print("nested tables:", len(cell.tables))


if __name__ == "__main__":
    main()
