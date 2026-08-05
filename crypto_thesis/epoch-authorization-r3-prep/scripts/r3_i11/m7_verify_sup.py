# -*- coding: utf-8 -*-
"""M7: verify citation tokens render as superscript runs in the DOCX."""
from __future__ import annotations

import re
import sys

from docx import Document


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    p = r"D:\Research\crypto_thesis\epoch-authorization-r3-prep\docs\midterm-report\m7\output\王威-专业学位研究生学位论文中期考评表-M7最终候选稿.docx"
    doc = Document(p)
    sup = 0
    non_sup = 0
    seen = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        if re.fullmatch(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]", r.text or ""):
                            if r.font.superscript:
                                sup += 1
                                if len(seen) < 6:
                                    seen.append(r.text)
                            else:
                                non_sup += 1
    print("superscript citation runs:", sup)
    print("non-superscript citation runs:", non_sup)
    print("sample:", seen)


if __name__ == "__main__":
    main()
