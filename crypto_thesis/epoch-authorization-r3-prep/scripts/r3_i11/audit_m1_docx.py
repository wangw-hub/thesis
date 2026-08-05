"""Verify the M1 midterm form DOCX content."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
DOCX = ROOT / "docs/midterm-report/output/王威-专业学位研究生学位论文中期考评表-候选稿.docx"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    d = Document(str(DOCX))
    print("tables:", len(d.tables))
    t0 = d.tables[0]
    print("r0:", t0.rows[0].cells[0].text.strip()[:70])
    r4 = t0.rows[4].cells[0].text
    print("progress len:", len(r4), "| starts:", r4[:50])
    for kw in ["图1", "图2", "图3", "15120", "145", "98.61", "I*", "C(P)", "CAP2", "Nonce", "Fail-Closed",
               "Kubo", "前瞻性撤销"]:
        print(f"  {kw}: {r4.count(kw)}")
    r6 = t0.rows[6].cells[0].text
    print("stage results:", r6[:150].replace("\n", " "))
    t1, t2, t3 = d.tables[1], d.tables[2], d.tables[3]
    print("problems len:", len(t1.rows[0].cells[0].text))
    print("solutions len:", len(t2.rows[0].cells[0].text))
    print("opinions r0 len:", len(t3.rows[0].cells[0].text), "| r7:", t3.rows[7].cells[0].text.strip()[:50])
    # cover fields
    cover = "\n".join(p.text for p in d.paragraphs[:14])
    for kw in ["王   威", "202422081113", "计算机技术", "高建彬", "王鹏", "2026    年  7  月  27"]:
        print(f"cover {kw!r}:", kw in cover)


if __name__ == "__main__":
    main()
