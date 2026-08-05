# -*- coding: utf-8 -*-
"""M1: build the midterm assessment form DOCX from the user's own template.

Base = 王威专业学位研究生学位论文中期考评表.docx (user's own, authoritative).
Rewrites 论文研究进展 / 存在问题和解决办法 with the current frozen research
content, corrects the opening date, embeds real figures, keeps 阶段性成果 and
the blank review-opinion rows.
"""
from __future__ import annotations

import os
import shutil
import stat
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
USER_FORM = Path(r"D:\Users\wangw\Desktop\中期和小论文\王威专业学位研究生学位论文中期考评表.docx")
OUT = ROOT / "docs/midterm-report/output"
OUT_DOCX = OUT / "王威-专业学位研究生学位论文中期考评表-候选稿.docx"
OUT_PDF = OUT / "王威-专业学位研究生学位论文中期考评表-候选稿.pdf"

FIG4 = Path(r"D:\Research\crypto_thesis\time-policy\figures\图4-1确定性时间策略编译流程.png")
FIG5 = Path(r"D:\Research\crypto_thesis\epoch-authorization\docs\thesis-drafts\research-content-2-final\figures\figure-5-1-design.png")
FIG6 = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep\experiments\r3\formal\figures\i12-final\fig-rq2-header-only-duration.png")


def set_run(r, east="宋体", latin="Times New Roman", size=None, bold=False):
    r.font.name = latin
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), east)
    r.bold = bold
    if size:
        r.font.size = Pt(size)


def clear_cell(cell) -> None:
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def add_para(cell, text: str, bold=False, size=12, align=None) -> None:
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(20)
    pf.line_spacing_rule = 4  # EXACTLY
    pf.space_after = Pt(0)
    if align:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, bold=bold, size=size)
    return p


def add_figure(cell, path: Path, caption: str) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    try:
        r.add_picture(str(path), width=Cm(12.5))
    except Exception:
        pass
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    set_run(cr, size=10.5)


def rebuild_cell(cell, blocks: list[tuple[str, str]]) -> None:
    """blocks: list of ('p', text) / ('h', text) / ('fig', path, caption) handled separately."""
    clear_cell(cell)
    for item in blocks:
        if item[0] == "fig":
            add_figure(cell, Path(item[1]), item[2])
        elif item[0] == "p":
            add_para(cell, item[1])
        elif item[0] == "h":
            add_para(cell, item[1], bold=True)


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    OUT.mkdir(parents=True, exist_ok=True)
    shutil.copy(USER_FORM, OUT_DOCX)
    os.chmod(OUT_DOCX, stat.S_IWRITE)
    doc = Document(str(OUT_DOCX))

    t0 = doc.tables[0]
    # 1. opening date correction 2026 -> 2025 (evidence: opening form 填表 2025-12-08)
    for cell in t0.rows[0].cells:
        txt = cell.text
        if "开题报告通过时间" in txt and "2026" in txt:
            new_txt = txt.replace("2026", "2025")
            for p in cell.paragraphs:
                for r in p.runs:
                    if "2026" in r.text:
                        r.text = r.text.replace("2026", "2025")
            if "2026" not in cell.text:
                print("opening date corrected ->", cell.text.strip())

    # 2. rebuild 论文研究进展 (table0 r4)
    progress_blocks: list[tuple[str, str]] = [
        ("h", "（1）研究背景与问题动因"),
        ("p", "随着医疗健康、政务协同、工业互联网和供应链金融等场景持续数字化，数据逐渐由单一机构内部资源转变为跨主体协同要素。区块链凭借不可篡改账本、智能合约和可追溯事件记录，为多方环境下的状态一致性与责任审计提供了技术基础，但区块链本身并不提供数据机密性：将业务数据直接写入链上会造成隐私暴露和高额存储成本，而仅将密文存放在链下、不约束密钥释放，又会使链上授权记录退化为事后审计附件。因此，本研究以“链下保存密文主体、链上维护公开授权状态、密码机制约束文件密钥释放”为基本架构，围绕状态控制与密码执行之间的闭合关系展开。"),
        ("p", "从时间维度看，面向数据共享的授权策略往往不是单一连续区间。在排班访问、周期授权、跨日窗口和节假日例外等业务中，时间条件通常由多个不连续、非对齐的窗口组成。不同输入序列虽然可能表达相同的允许时间集合，却会因乱序、重复、相交、嵌套或等价拆分而产生不同的字节表示和策略标识，破坏跨组件校验的一致性；若将每个最小时间单元直接写入策略，密文头部和策略处理成本又会随时间域快速膨胀。因此，第一个问题是如何为非连续时间约束建立确定、唯一、可复现的语义表示与策略摘要。"),
        ("p", "从状态维度看，仅有策略表示并不足以构成可信授权。用户即使持有满足策略的长期凭证，其访问仍可能因暂停、撤销、策略更新或密钥轮换而不再有效；若仅依赖本地或单一服务端校验，难以应对状态变更、跨实例重放以及链与合约上下文绑定。区块链能够提供确定性的公开状态与可审计的状态变迁，但如何把策略摘要、资源状态、用户密钥版本与授权时效锚定到链上，并使能力只能用于指定链与合约实例，是第二个需要回答的问题。在此基础上还存在第三个问题：授权状态变化与链下密文对象生命周期之间存在失配窗口，链上记录“用户已被撤销”并不自动意味着该用户无法继续获取链下密钥材料，撤销、文件更新、密钥轮换等事件必须传导到材料释放判定与密文对象版本上，才能构成完整治理闭环。"),
        ("p", "综上，本研究以“策略表示—可信授权执行—密文对象生命周期管理”为主线设置三项递进研究内容：其一解决非连续时间策略的确定性表示与编译，其二解决基于许可联盟链状态锚定的可信授权执行，其三解决授权状态变化后的版本化密文头部与前瞻性撤销闭环。三者共享半开区间时间语义与策略摘要，前一阶段输出的语义与状态作为后一阶段的输入，最终形成从策略编译、授权判定到材料释放与恢复的可验证闭合链路。"),
        ("h", "（2）研究目标与关键问题"),
        ("p", "本研究的总体目标是面向具有非连续时间约束和动态授权需求的区块链数据共享场景，建立一套安全边界清晰、协议接口闭合、能够通过原型与实验复核的可信共享方案。系统不追求把大文件或秘密计算迁移到链上，而是利用许可联盟链确定公开授权状态与请求顺序，利用能力结构与原子一次性机制实现跨实例一致的可信执行，并通过版本化密文头部把撤销与更新传导到链下材料释放。三个关键问题为：非连续时间约束如何获得确定、唯一、可复现的语义表示与策略摘要；确定的授权策略如何在许可联盟链环境中形成可信、可审计、抗重放的授权执行；授权状态变化后，链下密文对象、密钥材料与链上状态如何维持版本一致性并实现前瞻性撤销与故障恢复。"),
        ("h", "（3）研究内容与技术路线"),
        ("p", "研究内容一研究非连续时间策略的确定性规范化编译方法：以统一时区、最小粒度和半开区间为语义基础，以规范区间序列表达唯一语义，通过规范化、最大二进制层次覆盖、固定宽度规范编码与摘要计算，把任意等义输入映射为唯一策略标识，并分析语义保持性、规范唯一性、最大覆盖规范性与输出敏感复杂度。研究内容二研究基于许可联盟链状态锚定的可信授权执行机制：在真实多节点 Besu QBFT 许可联盟链上维护资源与用户状态，构造与链标识、合约地址、策略摘要、资源状态和用户密钥版本完整绑定的能力结构，利用数据库原子共享 Nonce 协调多个验证实例，并在依赖故障时保持 Fail-Closed。研究内容三研究版本化密文头部与前瞻性撤销闭环机制：以 Header、Body 与内容密钥的独立版本描述密文对象更新，通过链上 HeaderRegistry、数据库任务状态机与链下不可变对象存储形成可核验的闭合关系，撤销后立即停止后续材料释放，并在对象损坏等故障场景下提供可验证的恢复来源。"),
        ("h", "（4）研究内容一阶段性进展"),
        ("p", "设计层面明确了规范区间序列 I* 作为唯一语义表示：它是策略允许槽集合的极大连续分量分解，任何等义输入经规范化后得到同一序列；策略摘要绑定时间解释环境与 I*，不随执行层表示变化；层次覆盖 C(P) 被定位为可再生成的派生执行表示，NTP1 固定宽度大端编码保证摘要计算的确定性。实现与理论分析层面完成了确定性编译流程的原型（规范化、最大层次覆盖、规范编码与摘要计算），建立了语义保持性、规范序列唯一性、最大层次覆盖规范性等分析结论与 O(n log n + c) 的输出敏感复杂度刻画。"),
        ("p", "实验验证方面，正确性验证完成 81 项自动化测试（含性质测试与小域穷举），分支感知代码覆盖率为 98.61%，未发现偏离形式化定义的反例；性能方面完成 E1 正式实验，覆盖 168 个样本（E1-A 108、E1-B 36、E1-C 24），共 15120 条正式记录且全部有效，对时间槽枚举、普通规范区间列表与二进制层次覆盖三种表示进行比较。阶段性结果表明，层次覆盖相对普通规范区间列表没有普遍的存储或一维查询优势，高碎片策略下会明确退化；据此将规范区间保留为语义、摘要与普通匹配表示，把层次覆盖限定为派生执行结构与后续授权协议的可选接口，并如实报告这一负结果作为方法适用边界的组成部分。"),
        ("fig", str(FIG4), "图1 非连续时间策略确定性编译流程"),
        ("h", "（5）研究内容二阶段性进展"),
        ("p", "设计层面明确了单纯本地策略校验无法解决状态变更、跨实例重放与链/合约上下文绑定问题。系统在真实五节点 Besu QBFT 许可联盟链上部署 AuthorizationState 合约，分别维护资源状态与用户状态；构造 CAP2 能力结构，将链标识、合约地址、策略摘要、epoch、资源状态版本与用户密钥版本完整绑定；利用 PostgreSQL 唯一约束实现共享原子 Nonce 的一次性消费，设置两个相互独立的 Verifier 与数据库中断时的 Fail-Closed 行为。实现层面完成五节点联盟链部署、合约开发、能力签发与验证流程、共享 Nonce 服务以及依赖故障注入测试，并通过攻击回归、语义一致性、并发重放与状态竞争测试。"),
        ("p", "正式实验覆盖 108 个因素配置、324 个含 seed 配置、9720 个运行块、77760 条请求记录与 233280 条链读取记录；四种方法复用相同工作负载，以 2430 对运行块自然配对并执行 10000 次运行级 Bootstrap。阶段性结果表明，四种方法端到端中位时延均约为 196～199 ms，吞吐量中位数约为 17.78～17.93 请求/s，逐请求链读取占端到端时延的 98.66%～98.80%；缓存命中率的提高未稳定转化为端到端收益，C(P) 亦未表现出 Baseline-I 难以复制的协议能力或性能价值。基于这些结果，研究将贡献收敛到链上状态锚定、能力完整绑定、共享 Nonce、多验证实例一致性与 Fail-Closed，并据此完成技术路线的收窄与确认。"),
        ("fig", str(FIG5), "图2 许可授权执行正式实验因素与运行级配对结构"),
        ("h", "（6）研究内容三阶段性进展"),
        ("p", "设计层面明确了链上撤销记录与链下材料释放之间的失配问题，设计了版本化密文头部（Header/Body/CK 各自携带版本）、HEADER_ONLY 与 BODY_ROTATION 两类语义操作、基于链上复合状态与 Header 完整性的材料释放判定、链上 HeaderRegistry 提交记录、数据库任务状态机，以及 LocalObjectStore 与隔离 Kubo 副本组成的链下对象层，使授权状态、任务状态与链下不可变对象形成可验证的闭合关系。"),
        ("p", "正式实验在独立于 Pilot 的受控单节点环境中完成 29 个配置、35 个 warm-up 与 145 个 measured RUN，覆盖 E1 状态一致性与幂等性、E2 HEADER_ONLY 规模影响、E3 BODY_ROTATION 开销、E4 撤销窗口 Fail-Closed 与 E5 故障恢复。阶段性结果表明，状态一致性与幂等性检查全部通过，未观察到错误材料释放；撤销窗口内材料释放保持 Fail-Closed，Header 闭合后恢复释放；HEADER_ONLY 下接收者与受影响资源规模对端到端时延影响较小，BODY_ROTATION 在 8 MiB 规模下存在可观察的额外成本；Kubo 副本在特定本地对象损坏场景下提供可验证的恢复来源，而正常路径未观察到稳定的性能优势，因此被定位为故障恢复可用性机制。系统明确采用前瞻性撤销语义，不主张追回此前已合法获得的明文或旧密钥。"),
        ("fig", str(FIG6), "图3 HEADER_ONLY 操作端到端时延分布（阶段性实验结果）"),
        ("h", "（7）当前整体系统闭环与阶段性认识"),
        ("p", "截至目前，三项研究内容的核心技术路线与主要原型均已形成，关键实验取得阶段性结果，研究主线基本闭合。研究过程中形成了三次认识收敛：层次覆盖的存储优势假设未得到实验支持，被降级为派生执行表示；缓存与层次覆盖的端到端性能收益不稳定，贡献收敛到状态锚定、绑定、重放控制与故障闭合；Kubo 副本的正常路径性能优势不明显，被定位为恢复可用性机制。整体完成度可概括为：研究内容一核心方法与正式实验已完成，后续为论文级凝练与边界表述；研究内容二原型与正式实验已完成，后续为缓存与层次接口的协议级讨论；研究内容三原型与正式实验已完成，后续为多节点扩展与批量更新压力；论文整体进入整合、完善与定稿阶段。"),
    ]
    r4 = t0.rows[4].cells[0]
    instruction = "按照开题计划，填写开题以来学位论文工作的研究进展。视具体研究内容，可包括理论、计算、实验（或实证）等方面（可续页）"
    rebuild_cell(r4, progress_blocks)
    instr_p = r4.paragraphs[0].insert_paragraph_before()
    ir = instr_p.add_run(instruction)
    set_run(ir, size=12)

    # 3. rebuild 存在问题 (table1 r0) and 解决办法 (table1 r1)
    t1 = doc.tables[1]
    t2 = doc.tables[2]
    problems = [
        ("h", "1．未按开题计划完成的研究工作、存在的原理性/技术性难题及实验条件限制（可续页）"),
        ("p", "总体来看，当前研究已经形成“策略表示—可信授权执行—密文生命周期治理”的基本技术闭环，三项研究内容的核心算法与原型实验均已达到阶段性完成状态，但仍存在以下需要进一步解决的问题。"),
        ("p", "（1）三项研究内容的论文级整合与理论表述仍需深化。目前三项研究内容分别形成了较为完整的方法与实验结果，但仍需要把三者在统一框架下的接口关系、安全假设与结论边界整理为连续、自洽的学术论证；理论分析方面，复杂度刻画与正确性分析目前以性质论证和实验验证为主，尚未形成完整的端到端形式化安全归约，论文中需要明确区分“实验验证”与“形式化证明”的边界。"),
        ("p", "（2）实验的外部有效性受环境条件限制。研究内容三的正式实验在受控单节点环境中完成，尚不能评估多 Validator 共识对端到端时延的影响；研究内容二的五节点实验运行于共享物理主机的虚拟机之上，相关性能结论只能在冻结配置与环境范围内成立，不能外推到任意规模、公网环境或独立物理集群。"),
        ("p", "（3）相关工作与对比方案的覆盖仍需完善。文献方面已完成核心文献的真实性核验，但近五年许可链授权状态管理、跨链令牌绑定、版本化密文与前瞻撤销等主题的更广泛综述仍需补充；与属性基加密、门限解密等数据共享路线的对比关系也需在论文中进一步展开。"),
        ("p", "（4）部分机制的扩展性验证仍不充分。缓存机制与层次覆盖接口的协议级收益、批量密文头更新在更大文件数量下的开销、以及多验证实例在大并发下的表现，目前仅在冻结配置范围内得到验证，需要结合论文论证需要补充针对性实验或明确其作为未来工作的边界。"),
    ]
    rebuild_cell(t1.rows[0].cells[0], problems)

    solutions = [
        ("h", "2．针对上述问题的解决办法及下一步研究计划（可续页）"),
        ("p", "针对上述问题，后续将围绕论文整合与理论深化、实验边界补强、相关工作完善三条主线继续推进。"),
        ("p", "（1）针对论文级整合与理论表述问题，按照“统一语义—授权执行—密文生命周期”的主线重构论文论证结构，明确三项研究内容之间输入输出与接口关系；补充形式化符号、假设与结论边界的统一表述，严格区分实验验证与形式化证明，避免对安全性质作过强概括。"),
        ("p", "（2）针对实验外部有效性问题，如实保留现有结论的适用范围表述，并在条件允许时补充独立物理集群或多节点规模验证实验；无法在中期阶段完成的扩展实验，在论文中明确列为局限与未来工作。"),
        ("p", "（3）针对相关工作覆盖问题，基于已核验的核心文献补充近五年相关主题综述并完成与属性基加密、门限解密等路线的对比分析；全部新增文献按学校要求经真实来源核验后引入。"),
        ("p", "（4）针对扩展性验证问题，根据论文论证需要设计补充实验；若实验条件不支持，则以明确的边界表述纳入论文局限。"),
        ("p", "下一步具体研究计划（时间以实际中期考评与毕业安排为准）：阶段1（2026年9月—10月）完成三项研究内容的论文级整合与相关工作综述；阶段2（2026年11月—12月）深化理论表述与结论边界，按需补充针对性实验，完成全文图表、公式与参考文献规范化；阶段3（2027年1月—3月）完成学位论文全文定稿、格式审查与盲审准备，并根据导师与专家意见进行最后一轮修改。"),
    ]
    rebuild_cell(t2.rows[0].cells[0], solutions)

    doc.save(OUT_DOCX)
    print("saved:", OUT_DOCX)


if __name__ == "__main__":
    main()
