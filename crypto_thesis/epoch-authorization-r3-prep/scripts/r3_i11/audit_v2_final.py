"""Comprehensive V2 final audit: cover, sections, tables, refs, paragraphs."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
DOCX = ROOT / "docs/final-manuscript/output/THESIS-FORMAT-CANDIDATE-V2.docx"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(DOCX))
    text = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)

    print("=== cover fields ===")
    for kw in ["202422081113", "王威", "高建彬", "计算机技术", "计算机科学与工程学院", "关键技术研究及实现", "[待填写]"]:
        print(f"  {kw}: {text.count(kw)}")

    print("=== section headers/order ===")
    for kw in ["摘\u3000要", "ABSTRACT", "目\u3000录", "第一章 绪论", "致\u3000谢", "参考文献", "附录A 复现说明",
               "攻读硕士学位期间取得的成果", "第五章 链上状态驱动的可信授权执行机制"]:
        idx = text.find(kw)
        print(f"  {kw}: pos={idx}")

    print("=== table captions ===")
    caps = sorted(set(re.findall(r"表[45]-\d+", text)))
    print("  ", caps)
    print("=== figure captions ===")
    figs = sorted(set(re.findall(r"图[456]-\d+", text)))
    print("  ", figs)
    print("=== equation numbers ===")
    print("  ", sorted(set(re.findall(r"\([4-6]-\d+\)", text)))[:12])
    print("=== references sample ===")
    refs = [p.text for p in doc.paragraphs if re.match(r"^\[\d+\] ", p.text)]
    print("  count:", len(refs))
    for r in refs[:4]:
        print("  ", r[:150])
    print("=== manual breaks ===")
    breaks = text.count("\n")
    br_xml = doc.element.body.xml.count("<w:br")
    print("  <w:br> in body xml:", br_xml)
    # paragraph statistics (body paragraphs only, exclude headings/captions/refs)
    body_paras = []
    for p in doc.paragraphs:
        s = p.text.strip()
        if not s:
            continue
        if re.match(r"^(第[一二三四五六七]章|附录|参考文献|致|攻读|\d+\.\d)", s) or \
                re.match(r"^(图|表|算法)\s?\d", s) or re.match(r"^\[\d+\]", s) or \
                s.startswith(("摘\u3000要", "ABSTRACT", "目\u3000录", "关键词", "Keywords")):
            continue
        body_paras.append(s)
    lens = [len(s) for s in body_paras]
    one_sentence = sum(1 for s in body_paras if re.match(r"^[^。；;]+[。；;]$", s) and len(s) < 120)
    short = sum(1 for s in body_paras if len(s) < 60)
    print("  body paras:", len(body_paras), "avg chars:", round(sum(lens) / max(1, len(lens)), 1),
          "median:", sorted(lens)[len(lens) // 2], "max:", max(lens))
    print("  one-sentence paras:", one_sentence, "| <60 chars:", short)
    print("=== citations ===")
    head = text.split("参考文献", 1)[0]
    cited = sorted({int(m) for m in re.findall(r"\[(\d+)\]", head)})
    print("  cited:", cited)


if __name__ == "__main__":
    main()
