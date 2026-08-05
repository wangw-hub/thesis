# -*- coding: utf-8 -*-
"""M6: verify DOCX display equations (OMML) and inline math integrity."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
OUT = ROOT / "docs/midterm-report/m6/output/王威-专业学位研究生学位论文中期考评表-M6候选稿.docx"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    doc = Document(str(OUT))
    cell = doc.tables[0].rows[4].cells[0]
    paras = cell.paragraphs

    display = []
    inline = 0
    bad = 0
    for p in paras:
        omats = p._p.findall(f".//{M}oMath")
        if not omats:
            continue
        is_display = False
        for om in omats:
            pr = om.find(f".//{M}oMathPara")
            if pr is not None or om.getparent().tag == f"{M}oMathPara":
                is_display = True
        if is_display:
            # collect sibling text for number
            txt = p.text
            display.append(txt.strip())
        else:
            inline += 1
        # check for placeholder/garbage in OMML xml
        xml = etree.tostring(p._p, encoding="unicode")
        for tok in ("w:placeholder", "\ufffd", "MISSING"):
            if tok in xml:
                bad += 1

    print("display equation paragraphs:", len(display))
    for i, d in enumerate(display, 1):
        print(i, "|", d[:80])
    print("inline math runs:", inline)
    print("placeholder/garbage hits:", bad)


if __name__ == "__main__":
    main()
