# -*- coding: utf-8 -*-
"""M7: locate the 阶段性研究成果 heading paragraph in the built DOCX."""
from __future__ import annotations

import sys

from docx import Document


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    p = r"D:\Research\crypto_thesis\epoch-authorization-r3-prep\docs\midterm-report\m7\output\王威-专业学位研究生学位论文中期考评表-M7最终候选稿.docx"
    doc = Document(p)
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                if "阶段性研究成果" in cell.text:
                    print(f"table {ti} row {ri} col {ci}: {cell.text[:120]!r}")
                    for pi, para in enumerate(cell.paragraphs):
                        if "阶段性研究成果" in para.text:
                            print(f"   para {pi}: {para.text!r}")
                            kn = para.paragraph_format.keep_with_next
                            print("   keep_with_next:", kn)


if __name__ == "__main__":
    main()
