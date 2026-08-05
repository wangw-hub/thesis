# -*- coding: utf-8 -*-
"""M3: build the refined midterm DOCX from the user's original cover authority."""
from __future__ import annotations

import copy
import os
import re
import shutil
import stat
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree
from latex2mathml.converter import convert as latex_to_mathml


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
USER_DOCX = Path(r"D:\Users\wangw\Desktop\中期和小论文\王威专业学位研究生学位论文中期考评表.docx")
SRC = ROOT / "docs/midterm-report/m3/M3-MIDTERM-SOURCE.md"
OUT = ROOT / "docs/midterm-report/m3/output"
OUT_DOCX = OUT / "王威-专业学位研究生学位论文中期考评表-M3候选稿.docx"
FIGDIR = ROOT / "docs/midterm-report/m3/figures"
RC1_RUN = Path(r"D:\Research\crypto_thesis\time-policy\experiments\runs\e1_20260727_ec8b193_r3\figures")
RC2_FIG = Path(r"D:\Research\crypto_thesis\epoch-authorization\docs\thesis-drafts\research-content-2-final\figures")
RC3_FIG = ROOT / "experiments/r3/formal/figures/i12-final"
MML_XSL = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"


# ---------------------------------------------------------------------------
# OMML
# ---------------------------------------------------------------------------

_XSLT = None


def _transformer():
    global _XSLT
    if _XSLT is None:
        _XSLT = etree.XSLT(etree.parse(MML_XSL))
    return _XSLT


def latex_to_omml(latex: str) -> etree._Element:
    mathml = latex_to_mathml(latex, display="block")
    tree = etree.fromstring(mathml.encode("utf-8"))
    res = _transformer()(tree)
    root = res.getroot()
    ns = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
    om = root.find(".//m:oMath", ns)
    if om is None:
        om = root
    return copy.deepcopy(om)


def add_omml(p, latex: str) -> None:
    p._p.append(latex_to_omml(latex))


# ---------------------------------------------------------------------------
# Text helpers
# ---------------------------------------------------------------------------

TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\*[^*\s][^*]*?\*|`[^`]+?`|\\\(.*?\\\)|\\\[.*?\\\]|\[\d+(?:,\s*\d+)*\]|\[\d+\])", re.S)


def set_run(run, bold=False, italic=False, code=False, sup=False, east="宋体", latin="Times New Roman", size=None):
    if code:
        latin = "Consolas"
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), east)
    run.bold = bold
    run.italic = italic
    if sup:
        run.font.superscript = True
    if size:
        run.font.size = Pt(size)


def render_inline(p, text: str, size: float = 12, east="宋体"):
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); set_run(r, bold=True, size=size, east=east)
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); set_run(r, code=True, size=size - 1, east=east)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = p.add_run(tok[1:-1]); set_run(r, italic=True, size=size, east=east)
        elif tok.startswith("\\(") and tok.endswith("\\)"):
            add_omml(p, tok[2:-2])
        elif tok.startswith("\\[") and tok.endswith("\\]"):
            add_omml(p, tok[2:-2])
        elif re.fullmatch(r"\[\d+(?:,\s*\d+)*\]", tok):
            r = p.add_run(tok); set_run(r, sup=True, size=size, east=east)
        else:
            r = p.add_run(tok); set_run(r, size=size, east=east)


def para_props(p, align=None, first_chars=0, hanging_chars=0, line_pt=20,
               space_before=0, space_after=0, left_cm=0.0):
    pf = p.paragraph_format
    pf.line_spacing = Pt(line_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    if first_chars:
        ind.set(qn("w:firstLineChars"), str(first_chars * 100))
    if hanging_chars:
        ind.set(qn("w:hangingChars"), str(hanging_chars * 100))
    if left_cm:
        ind.set(qn("w:left"), str(int(left_cm * 567)))


def clear_cell(cell):
    for t in list(cell.tables):
        t._tbl.getparent().remove(t._tbl)
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def add_para(cell, text, bold=False, size=12, first_chars=2, east="宋体", align=None,
             line_pt=20, hang=0):
    p = cell.add_paragraph()
    para_props(p, align=align, first_chars=first_chars, hanging_chars=hang, line_pt=line_pt)
    if bold:
        r = p.add_run(text)
        set_run(r, bold=True, size=size, east=east)
    else:
        render_inline(p, text, size=size, east=east)
    return p


# ---------------------------------------------------------------------------
# Figure / table / algorithm registries
# ---------------------------------------------------------------------------

FIGURE_MAP = [
    ("图1 论文总体技术路线与三项研究内容递进关系", FIGDIR / "m3-fig1-technical-route.png"),
    ("图2 非连续时间策略确定性编译流程", FIGDIR / "m3-fig2-compile-flow.png"),
    ("图3 语义主表示—摘要—派生执行IR关系", FIGDIR / "m3-fig3-ir-relation.png"),
    ("图4 匹配查询中位时延（E1-A 正式实验结果）", FIGDIR / "m3-exp-fig10-match.png"),
    ("图5 三种表示的逻辑规模比较（E1-A 正式实验结果）", FIGDIR / "m3-exp-fig8-rep-size.png"),
    ("图6 表示的压缩比与适用边界（E1-A 正式实验结果）", FIGDIR / "m3-exp-fig11-boundary.png"),
    ("图7 许可联盟链可信授权系统总体架构", FIGDIR / "m3-fig4-arch.png"),
    ("图8 CAP2 签发与验证双泳道流程", FIGDIR / "m3-fig5-cap2-swimlane.png"),
    ("图9 并发度对端到端时延的影响（RC2）", FIGDIR / "m3-exp-fig15-concurrency.png"),
    ("图10 四种方法的运行级端到端时延分布（RC2 正式实验结果）", FIGDIR / "m3-exp-fig12-run-latency.png"),
    ("图11 请求局部性与缓存的影响（RC2）", FIGDIR / "m3-exp-fig17-locality.png"),
    ("图12 端到端时延的阶段占比（RC2 中位数）", FIGDIR / "m3-exp-fig13-stage-share.png"),
    ("图13 自然配对比较与运行级 Bootstrap 置信区间（RC2）", FIGDIR / "m3-exp-fig14-paired.png"),
    ("图14 碎片率对匹配时延的影响（RC2）", FIGDIR / "m3-exp-fig16-fragmentation.png"),
    ("图15 版本化密文对象结构（Header/Body/CK）", FIGDIR / "m3-fig6-object-structure.png"),
    ("图16 链上可信状态—控制协调—链下密文对象三层闭环架构", FIGDIR / "m3-fig7-three-layer-closure.png"),
    ("图17 HEADER_ONLY 操作端到端时延（RC3 正式实验结果）", FIGDIR / "m3-exp-fig18-header-only.png"),
    ("图18 BODY_ROTATION 操作端到端时延（RC3 正式实验结果）", FIGDIR / "m3-exp-fig19-body-rotation.png"),
    ("图19 LOCAL_ONLY 与 KUBO_REPLICA 恢复时延对比（RC3 E5）", FIGDIR / "m3-exp-fig20-recovery.png"),
]


def fig_for(marker):
    for cap, path in FIGURE_MAP:
        key = cap.split(" ")[0] + " " + cap.split(" ", 1)[1].split("（")[0].strip()
        if marker.startswith("[方法图：") and key in marker:
            return path, cap
    return None, None


TABLE_DATA = {
    "三种表示的理论与实现特征": (
        ["表示", "构造复杂度", "查询复杂度", "E1-A样本逻辑字节中位数", "E1-A查询中位数/ns", "主要用途"],
        [["时间槽枚举", "O(A)", "期望 O(1)", "24000", "350.4", "小域、高频查询"],
         ["规范区间列表", "O(n log n)", "O(log k)", "2808", "561.0", "一维存储与匹配"],
         ["层次覆盖", "O(n log n + c)", "O(log U)", "3664", "1984.7", "层次授权接口"]]),
    "系统安全目标、机制与证据及结论边界": (
        ["目标", "机制与证据", "结论边界"],
        [["S1 状态锚定与策略绑定", "资源/用户状态上链，policyDigest 绑定", "不证明链本身绝对可信"],
         ["S2 能力完整绑定", "CAP2 绑定链/合约/策略/版本，篡改测试拒绝", "不抵抗合法角色恶意使用"],
         ["S3 跨实例重放控制", "共享原子 Nonce，并发 50/100/500 均仅一次成功", "限于冻结命名空间与事务语义"],
         ["S4 依赖故障闭合", "RPC/数据库故障时拒绝授权", "不构成形式化安全证明"]]),
    "四种方法运行级总体统计": (
        ["方法", "运行数", "中位时延/ms", "均值/ms", "吞吐量中位数/(请求/s)", "缓存命中率中位数", "链读取占比/%"],
        [["B0", "2430", "196.128", "209.714", "17.926", "0", "98.796"],
         ["B1", "2430", "196.624", "210.462", "17.906", "0.125", "98.757"],
         ["C0", "2430", "198.167", "211.739", "17.783", "0", "98.702"],
         ["C1", "2430", "198.601", "212.114", "17.786", "0.125", "98.670"]]),
    "四种自然配对比较及运行级 Bootstrap 置信区间": (
        ["比较", "配对数", "中位差/ms", "均值差/ms", "均值差 95% CI/ms", "改善比例", "退化比例"],
        [["B1-B0", "2430", "+0.390", "+1.688", "[-0.220, 3.539]", "43.70%", "47.33%"],
         ["C1-C0", "2430", "+0.176", "+0.318", "[-1.630, 2.267]", "44.32%", "46.71%"],
         ["C0-B0", "2430", "+0.704", "+1.416", "[-0.886, 3.718]", "-", "-"],
         ["C1-B1", "2430", "+0.408", "+1.046", "[-0.717, 2.809]", "-", "-"]]),
    "正式实验配置与运行汇总": (
        ["实验", "路径/变量", "有效运行数", "核心检查", "结果"],
        [["E1", "INITIAL/BODY_ROTATION/REVOCATION/RESTORE", "20", "状态一致与幂等", "全部通过，错误材料释放 0"],
         ["E2", "HEADER_ONLY/BODY_ROTATION 成本", "30", "版本关系", "HEADER_ONLY 接近固定成本"],
         ["E3", "撤销窗口判定", "15", "Fail-Closed", "pending 拒绝、闭合后允许"],
         ["E4", "pending/闭合路径", "10", "材料释放判定", "拒绝/允许边界可追踪"],
         ["E5", "来源×故障类别", "40", "恢复判定", "损坏场景副本可恢复，其余 Fail-Closed"]]),
    "E5 恢复结果与时长汇总": (
        ["故障", "对象来源", "有效(n)", "恢复判定", "修复动作", "时长中位数/ms"],
        [["对象损坏", "LOCAL_ONLY", "5", "FAIL_CLOSED", "0", "3112.2"],
         ["对象损坏", "KUBO_REPLICA", "5", "CONSISTENT", "1", "3129.6"],
         ["CID 不一致", "LOCAL_ONLY", "5", "FAIL_CLOSED", "0", "-"],
         ["CID 不一致", "KUBO_REPLICA", "5", "FAIL_CLOSED", "0", "-"],
         ["同时缺失", "LOCAL_ONLY", "5", "FAIL_CLOSED", "0", "-"],
         ["同时缺失", "KUBO_REPLICA", "5", "FAIL_CLOSED", "0", "-"]]),
    "正式实验因素设计汇总": (
        ["因素/规模", "水平或取值", "配置数", "说明"],
        [["碎片率", "0 / 0.5 / 1", "3", "策略碎片程度"],
         ["请求局部性", "均匀 / 区间热点 / 节点热点", "3", "工作负载生成器"],
         ["并发度", "1 / 4 / 16", "3", "运行级并发"],
         ["随机种子", "固定 3 个 seed", "3", "含 seed 配置 324"],
         ["重复", "每含 seed 配置 30 次正式重复", "-", "9720 个运行块"],
         ["请求/链读取", "77760 请求 / 233280 链读取", "-", "每请求三次真实链读取"],
         ["配对统计", "每组比较 2430 对运行块", "-", "10000 次运行级 Bootstrap"]]),
}


def add_table(cell, header, rows, caption=None, width_cm=15.5, font=8.5):
    if caption:
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        set_run(cr, size=10.5)
        para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=16)
    ncols = max(len(header), max((len(r) for r in rows), default=1))
    t = cell.add_table(rows=1 + len(rows), cols=ncols)
    t.style = None
    tbl = t._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    total_twips = 9628
    widths = [total_twips // ncols] * ncols
    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
            gc.set(qn("w:w"), str(w))
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    for ri, row in enumerate([header] + rows):
        for ci in range(ncols):
            c = t.cell(ri, ci)
            tcPr = c._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[ci] if ci < len(widths) else widths[-1]))
            tcW.set(qn("w:type"), "dxa")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    for edge in ("insideH",):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)
    for ci, h in enumerate(header):
        c = t.cell(0, ci)
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = Pt(14)
        r = p.add_run(h)
        set_run(r, size=font, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows, 1):
        for ci in range(ncols):
            c = t.cell(ri, ci)
            p = c.paragraphs[0]
            p.paragraph_format.line_spacing = Pt(14)
            r = p.add_run(row[ci] if ci < len(row) else "")
            set_run(r, size=font)
            if ri == len(rows):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def add_algo(cell, text):
    lines = text.splitlines()
    p = cell.add_paragraph()
    para_props(p, line_pt=14)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "404040")
        pbdr.append(el)
    ppr.append(pbdr)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        set_run(r, size=9, latin="Consolas")
        if i < len(lines) - 1:
            r.add_break()


EQ_COUNT = [0]


def add_equation(cell, latex):
    EQ_COUNT[0] += 1
    n = EQ_COUNT[0]
    p = cell.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    te = OxmlElement("w:tab")
    te.set(qn("w:val"), "right")
    te.set(qn("w:pos"), str(int(16.5 * 567)))
    tabs.append(te)
    para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=18)
    add_omml(p, latex)
    r = p.add_run(f"\t({n})")
    set_run(r, size=10.5)
    return n


def add_figure(cell, path, caption, width_cm=12.5):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_props(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=14)
    run = p.add_run()
    try:
        run.add_picture(str(path), width=Cm(width_cm))
    except Exception as exc:
        print("figure missing:", path, exc)
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    set_run(cr, size=10.5)
    para_props(cap, align=WD_ALIGN_PARAGRAPH.CENTER, line_pt=16)


# ---------------------------------------------------------------------------
# Block parsing
# ---------------------------------------------------------------------------

def build_blocks(lines):
    blocks = []
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith("[方法图："):
            path, cap = fig_for(s)
            blocks.append(("fig", path, cap if cap else s[5:-1]))
            i += 1
            continue
        if s.startswith("[表："):
            key = s[3:-1]
            for k, (h, rows) in TABLE_DATA.items():
                if k in key:
                    blocks.append(("table", h, rows))
                    break
            i += 1
            continue
        if s.startswith("[算法框："):
            blocks.append(("algo", s[5:-1]))
            i += 1
            continue
        if s.startswith("[公式："):
            blocks.append(("eq", s[4:-1]))
            i += 1
            continue
        if s.startswith("**（") and "）" in s:
            blocks.append(("h", re.sub(r"[*#|`]", "", s)))
            i += 1
            continue
        if s.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip().strip("*") for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            data = [r for r in rows if not all(re.fullmatch(r":?-{3,}:?", c) for c in r)]
            if data:
                blocks.append(("table", data[0], data[1:]))
            continue
        if s.startswith("### "):
            blocks.append(("h", s[4:].strip()))
            i += 1
            continue
        blocks.append(("p", s))
        i += 1
    return blocks


def fill_cell_blocks(cell, blocks, body_first=2, body_size=12):
    for blk in blocks:
        if blk[0] == "h":
            add_para(cell, blk[1], bold=True, first_chars=0, size=body_size)
        elif blk[0] == "p":
            add_para(cell, blk[1], first_chars=body_first, size=body_size)
        elif blk[0] == "fig":
            add_figure(cell, blk[1], blk[2])
        elif blk[0] == "table":
            add_table(cell, blk[1], blk[2])
        elif blk[0] == "algo":
            add_algo(cell, blk[1])
        elif blk[0] == "eq":
            add_equation(cell, blk[1])


def parse_refs(text):
    m = re.search(r"### 参考文献\n(.*?)(?=### 4．阶段性研究成果)", text, re.S)
    refs = []
    if m:
        for ln in m.group(1).splitlines():
            s = ln.strip()
            if s and s[0] == "[" and "]" in s:
                refs.append(s)
    return refs


def main():
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    OUT.mkdir(parents=True, exist_ok=True)
    shutil.copy(USER_DOCX, OUT_DOCX)
    os.chmod(OUT_DOCX, stat.S_IWRITE)
    doc = Document(str(OUT_DOCX))

    # cover: keep the user's original layout, fix the checked boxes only
    for para in doc.paragraphs:
        t = para.text
        if "攻读学位级别" in t:
            for r in para.runs:
                if r.text == "□":
                    r.text = "□"
                elif r.text == "硕士":
                    r.text = "☑硕士"

    draft = SRC.read_text(encoding="utf-8")
    lines = draft.splitlines()
    a = next(i for i, ln in enumerate(lines) if ln.startswith("**（1）研究背景"))
    b = next(i for i, ln in enumerate(lines) if ln.strip().startswith("### 4．阶段性研究成果"))
    blocks = build_blocks(lines[a:b])

    t0 = doc.tables[0]
    for cell in t0.rows[0].cells:
        if "开题报告通过时间" in cell.text:
            for p in cell.paragraphs:
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
                r = p.add_run("1．开题报告通过时间：2025年12月24日（以开题记录为准）")
                set_run(r, size=12)
                break
            break
    for cell in t0.rows[2].cells:
        if "学分要求" in cell.text:
            for p in cell.paragraphs:
                for r in p.runs:
                    if "□是" in r.text:
                        r.text = r.text.replace("□是", "☑是")

    # research progress
    r4 = t0.rows[4].cells[0]
    clear_cell(r4)
    instr_p = r4.add_paragraph()
    ir = instr_p.add_run("按照开题计划，填写开题以来学位论文工作的研究进展。视具体研究内容，可包括理论、计算、实验（或实证）等方面（可续页）")
    set_run(ir, size=12)
    fill_cell_blocks(r4, blocks)

    # references appended at the end of the progress cell
    refs = parse_refs(draft)
    if refs:
        add_para(r4, "参考文献", bold=True, first_chars=0)
        for ref in refs:
            add_para(r4, ref, first_chars=0, size=10.5, hang=2)

    # stage results
    r6 = t0.rows[6].cells[0]
    sres = ["按《研究生学位论文撰写格式规范》的格式要求分类填写与学位论文相关的阶段性研究成果（可续页）",
            "[1] 王威, 夏琦, 高建彬, 夏虎. 面向链上数据的双重绑定门限解封装方案[J]. 软件学报（拟投稿）.",
            "[2] [王威]. 一种非连续时间访问策略的压缩方法及系统: 中国, [P].（拟申请）.",
            "[3] [王威]. 一种基于属性与链上请求双重绑定的数据共享方法及系统: 中国, [P].（拟申请）.",
            "另：三套可复现原型与冻结实验数据集（时间策略编译、许可链授权执行、版本化密文头部与撤销恢复）。"]
    clear_cell(r6)
    for line in sres:
        add_para(r6, line, bold=line.startswith("["), first_chars=0)

    # problems & solutions (tables 1 and 2)
    c = draft.find("总体来看，当前研究已经形成")
    d = draft.find("针对上述问题，后续将围绕")
    e = draft.find("## 三、中期考评审查意见")
    t1 = doc.tables[1]
    clear_cell(t1.rows[0].cells[0])
    add_para(t1.rows[0].cells[0], "1．未按开题计划完成的研究工作，研究工作存在的原理性、技术性难题以及在实验条件等方面的限制（可续页）", bold=True, first_chars=0)
    for para in draft[c:d].splitlines():
        s = para.strip()
        if not s or s.startswith("## ") or s.startswith("### "):
            continue
        add_para(t1.rows[0].cells[0], s)
    t2 = doc.tables[2]
    clear_cell(t2.rows[0].cells[0])
    add_para(t2.rows[0].cells[0], "2．针对上述问题采取何种解决办法，对学位论文的研究内容及所采取的理论方法、技术路线和实施方案的进一步调整，以及下一步研究计划（可续页）", bold=True, first_chars=0)
    for para in draft[d:e].splitlines():
        s = para.strip()
        if not s or s.startswith("## ") or s.startswith("### "):
            continue
        add_para(t2.rows[0].cells[0], s)

    doc.save(OUT_DOCX)
    print("saved:", OUT_DOCX)
    print("equations numbered:", EQ_COUNT[0])
    print("figures:", len([x for x in blocks if x[0] == "fig"]))
    print("tables:", len([x for x in blocks if x[0] == "table"]))
    print("algorithms:", len([x for x in blocks if x[0] == "algo"]))
    print("refs:", len(refs))


if __name__ == "__main__":
    main()
