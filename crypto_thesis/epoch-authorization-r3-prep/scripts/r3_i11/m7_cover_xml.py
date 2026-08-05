# -*- coding: utf-8 -*-
"""M7: inspect cover runs (checks) and blank page 2 origin."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
TEMPLATE = Path(r"D:\Users\wangw\Desktop\中期和小论文\王威专业学位研究生学位论文中期考评表.docx")
BLANK = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\附件2：专业学位研究生学位论文中期考评表-2023版.docx")
M6_DOCX = ROOT / "docs/midterm-report/m6/output/王威-专业学位研究生学位论文中期考评表-M6候选稿.docx"


def dump_runs(path: Path, label: str) -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print("=" * 20, label)
    doc = Document(str(path))
    for i, p in enumerate(doc.paragraphs[:16]):
        if not p.text.strip():
            continue
        runs = [(r.text, r.font.name) for r in p.runs]
        print(f"p{i}: text={p.text[:80]!r}")
        print("   runs:", runs[:8])


def main() -> None:
    dump_runs(TEMPLATE, "USER TEMPLATE")
    dump_runs(M6_DOCX, "M6 DOCX")
    if BLANK.exists():
        dump_runs(BLANK, "SCHOOL BLANK 2023")
        d = Document(str(BLANK))
        print("blank template paragraphs count:", len(d.paragraphs))
        print("blank template tables:", len(d.tables))
    else:
        print("school blank template not found at expected path")


if __name__ == "__main__":
    main()
