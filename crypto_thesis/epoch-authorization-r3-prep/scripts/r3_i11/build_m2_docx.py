# -*- coding: utf-8 -*-
"""M2: build the full midterm form DOCX from the official blank template."""
from __future__ import annotations

import os
import re
import shutil
import stat
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\Research\crypto_thesis\epoch-authorization-r3-prep")
TEMPLATE = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\附件2：专业学位研究生学位论文中期考评表-2023版.docx")
DRAFT = ROOT / "docs/midterm-report/m2/MIDTERM-REPORT-M2-FULL-DRAFT.md"
OUT = ROOT / "docs/midterm-report/m2/output"
OUT_DOCX = OUT / "王威-专业学位研究生学位论文中期考评表-M2候选稿.docx"
FIGDIR = ROOT / "docs/midterm-report/m2/figures"
RC1_FIG = Path(r"D:\Research\crypto_thesis\time-policy\figures")
RC1_RUN = Path(r"D:\Research\crypto_thesis\time-policy\experiments\runs\e1_20260727_ec8b193_r3\figures")
RC2_FIG = Path(r"D:\Research\crypto_thesis\epoch-authorization\docs\thesis-drafts\research-content-2-final\figures")
RC3_FIG = ROOT / "experiments/r3/formal/figures/i12-final"


def set_run(r, east="宋体", latin="Times New Roman", size=None, bold=False):
    r.font.name = latin
    rpr = r._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    rf.set(qn("w:ascii"), latin)
    rf.set(qn("w:hAnsi"), latin)
    rf.set(qn("w:eastAsia"), east)
    r.bold = bold
    if size:
        r.font.size = Pt(size)


def clear_cell(cell):
    for p in list(cell.paragraphs):
        p._p.getparent().remove(p._p)


def add_para(cell, text, bold=False, size=12):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(20)
    pf.line_spacing_rule = 4
    pf.space_after = Pt(0)
    r = p.add_run(text)
    set_run(r, bold=bold, size=size)
    return p


def add_figure(cell, path, caption, width_cm=12.5):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(path), width=Cm(width_cm))
    except Exception:
        pass
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    set_run(cr, size=10.5)


def add_table(cell, header, rows):
    ncols = max(len(header), max((len(r) for r in rows), default=1))
    t = cell.add_table(rows=1 + len(rows), cols=ncols)
    for ci, h in enumerate(header):
        c = t.cell(0, ci)
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = Pt(14)
        r = p.add_run(h)
        set_run(r, size=9, bold=True)
    for ri, row in enumerate(rows, 1):
        for ci in range(ncols):
            c = t.cell(ri, ci)
            p = c.paragraphs[0]
            p.paragraph_format.line_spacing = Pt(14)
            r = p.add_run(row[ci] if ci < len(row) else "")
            set_run(r, size=9)


def add_algo(cell, text):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = Pt(14)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "808080")
        pbdr.append(el)
    ppr.append(pbdr)
    for line in text.splitlines():
        r = p.add_run(line)
        set_run(r, size=9, latin="Consolas")
        r.add_break()


FIGURE_MAP = [
    ("非连续时间策略确定性编译流程", RC1_FIG / "图4-1确定性时间策略编译流程.png", "图1 非连续时间策略确定性编译流程"),
    ("三种表示规模比较与匹配延迟", RC1_RUN / "figure_4_2_representation_size.png", "图2 三种表示规模比较（阶段实验结果）"),
    ("CAP2 能力签发与验证流程示意", FIGDIR / "schematic-cap2-flow.png", "图3 CAP2 能力签发与验证流程示意"),
    ("正式实验因素与运行级配对结构", RC2_FIG / "figure-5-1-design.png", "图4 正式实验因素与运行级配对结构"),
    ("四种方法端到端运行级时延分布", RC2_FIG / "figure-5-2-run-latency.png", "图5 四种方法端到端运行级时延分布"),
    ("闭环架构", FIGDIR / "schematic-closure-arch.png", "图6 链上授权状态、任务状态与链下对象闭环架构"),
    ("HEADER_ONLY 与 BODY_ROTATION 操作端到端时延分布", None, "图7 HEADER_ONLY 与 BODY_ROTATION 操作端到端时延分布"),
    ("LOCAL_ONLY 与 KUBO_REPLICA 恢复运行端到端时延对比", RC3_FIG / "fig-rq5-recovery-local-kubo.png", "图8 LOCAL_ONLY 与 KUBO_REPLICA 恢复运行端到端时延对比"),
]


def fig_for(marker):
    for kw, path, cap in FIGURE_MAP:
        if kw in marker:
            return path, cap
    return None, None


TABLE_DATA = {
    "三种表示的理论与实现特征": (
        ["表示", "构造复杂度", "查询复杂度", "E1-A样本逻辑字节中位数", "E1-A查询中位数/ns", "主要用途"],
        [["时间槽枚举", "O(A)", "期望 O(1)", "24000", "350.4", "小域、高频查询"],
         ["规范区间列表", "O(n log n)", "O(log k)", "2808", "561.0", "一维存储与匹配"],
         ["层次覆盖", "O(n log n + c)", "O(log U)", "3664", "1984.7", "层次授权接口"]]),
    "E2正确性验证汇总": (
        ["验证项目", "数量/结果", "失败或反例"],
        [["自动化测试(pytest)", "81 项通过", "0"],
         ["性质测试与穷举", "覆盖语义一致性、幂等性、置换不变性、摘要一致性、覆盖结构", "0"],
         ["分支感知代码覆盖率", "98.61%", "不适用"]]),
    "安全目标、机制与证据及结论边界": (
        ["目标", "机制与证据", "结论边界"],
        [["S1 状态锚定与策略绑定", "资源/用户状态上链，policyDigest 绑定", "不证明链本身绝对可信"],
         ["S2 能力完整绑定", "CAP2 绑定链/合约/策略/版本，篡改测试拒绝", "不抵抗合法角色恶意使用"],
         ["S3 跨实例重放控制", "共享原子 Nonce，并发 50/100/500 均仅一次成功", "限于冻结命名空间与事务语义"],
         ["S4 依赖故障闭合", "RPC/数据库故障时拒绝授权", "不构成形式化安全证明"]]),
    "四种方法运行级总体统计": (
        ["方法", "运行数", "中位时延/ms", "均值/ms", "吞吐量中位数/(请求/s)", "缓存命中率中位数", "链读取占比/%"],
        [["B0", "2430", "196.128", "209.714", "17.926", "0", "98.796"],
         ["B1", "2430", "196.624", "210.462", "17.906", "0.125", "98.757"],
         ["C0", "2430", "198.167", "211.739", "17.783", "0", "98.702"],
         ["C1", "2430", "198.601", "212.114", "17.786", "0.125", "98.670"]]),
    "四种自然配对比较及运行级 Bootstrap 置信区间": (
        ["比较", "配对数", "中位差/ms", "均值差/ms", "均值差 95% CI/ms", "改善比例", "退化比例"],
        [["B1-B0", "2430", "+0.390", "+1.688", "[-0.220, 3.539]", "43.70%", "47.33%"],
         ["C1-C0", "2430", "+0.176", "+0.318", "[-1.630, 2.267]", "44.32%", "46.71%"],
         ["C0-B0", "2430", "+0.704", "+1.416", "[-0.886, 3.718]", "-", "-"],
         ["C1-B1", "2430", "+0.408", "+1.046", "[-0.717, 2.809]", "-", "-"]]),
    "E5 恢复结果与时长汇总": (
        ["故障", "对象来源", "有效(n)", "恢复判定", "修复动作", "时长中位数/ms"],
        [["对象损坏", "LOCAL_ONLY", "5", "FAIL_CLOSED", "0", "3112.2"],
         ["对象损坏", "KUBO_REPLICA", "5", "CONSISTENT", "1", "3129.6"],
         ["CID 不一致", "LOCAL_ONLY", "5", "FAIL_CLOSED", "0", "-"],
         ["CID 不一致", "KUBO_REPLICA", "5", "FAIL_CLOSED", "0", "-"],
         ["同时缺失", "LOCAL_ONLY", "5", "FAIL_CLOSED", "0", "-"],
         ["同时缺失", "KUBO_REPLICA", "5", "FAIL_CLOSED", "0", "-"]]),
}


ALGORITHMS = {
    "POLICY-COMPILE": "算法1 POLICY-COMPILE\n输入：已离散化区间序列 P，时区感知起点 t0，粒度 Δ，槽总数 U\n输出：规范区间 I*，层次覆盖 C，规范字节 B，摘要 pd\n1: 校验 t0 含时区、Δ>0、U>0\n2: originUTC ← UTC(t0)\n3: I* ← NORMALIZE(P, U)\n4: C ← COVER-POLICY(I*, U)\n5: B ← NTP1-SERIALIZE(originUTC, Δ, U, I*)\n6: pd ← SHA256(B)\n7: 返回 (I*, C, B, pd)",
    "Verifier": "Verifier 验证流程\n1: 解析规范编码；失败返回 MALFORMED_TOKEN\n2: 验证签名；失败返回 INVALID_SIGNATURE\n3: 读取确认链状态；失败返回 SYSTEM_STATE_UNAVAILABLE\n4: 检查资源/用户/policyDigest/epoch/链与合约绑定/版本/时间窗口\n5: 重新执行 I* 时间策略检查\n6: 原子消费共享 Nonce；冲突返回 NONCE_REPLAY\n7: 全部通过且消费成功时返回 ACCEPT",
    "恢复协调": "恢复协调流程\n1: 读取候选对象（本地或隔离副本）\n2: SHA-256 摘要验证；不匹配返回不可恢复\n3: 结构验证（Header/Body 格式与版本关系）\n4: 原子恢复至 LocalObjectStore\n5: 记录修复来源与修复数量\n6: 形成一致状态或 Fail-Closed 结果",
}


def strip_md(text):
    return re.sub(r"[*#|`]", "", text).strip()


def build_progress_blocks(lines):
    blocks = []
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith("[图："):
            path, cap = fig_for(s)
            if path is None:
                blocks.append(("fig2", RC3_FIG / "fig-rq2-header-only-duration.png", RC3_FIG / "fig-rq3-body-rotation-duration.png", cap))
            else:
                blocks.append(("fig", path, cap))
            i += 1
            continue
        if s.startswith("[表："):
            key = s[3:-1]
            for k, (h, rows) in TABLE_DATA.items():
                if k in key:
                    blocks.append(("table", h, rows))
                    break
            i += 1
            continue
        if s.startswith("[算法："):
            key = s[3:-1]
            for k, txt in ALGORITHMS.items():
                if k in key:
                    blocks.append(("algo", txt))
                    break
            i += 1
            continue
        if s.startswith("**（") and "）" in s:
            blocks.append(("h", strip_md(s)))
            i += 1
            continue
        if s.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip().strip("*") for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            data = [r for r in rows if not all(re.fullmatch(r":?-{3,}:?", c) for c in r)]
            if data:
                blocks.append(("table", data[0], data[1:]))
            continue
        blocks.append(("p", s))
        i += 1
    return blocks


def fill_cover(doc):
    label_map = {
        "攻读学位级别": None, "培养方式": None,
        "专业学位类别或领域": "计算机技术",
        "学院": "计算机科学与工程学院（网络空间安全学院）",
        "学号": "202422081113",
        "姓名": "王  威",
        "论文题目": "面向非连续时间约束的区块链数据共享关键技术研究及实现",
        "校内指导教师": "高建彬",
        "校外指导教师": "王鹏",
    }
    for p in doc.paragraphs:
        t = p.text
        if "：" not in t:
            continue
        label = re.sub(r"\s+", "", t.split("：")[0])
        if label in ("攻读学位级别", "培养方式"):
            for r in p.runs:
                if "□硕士" in r.text:
                    r.text = r.text.replace("□硕士", "☑硕士")
                if "□全日制" in r.text:
                    r.text = r.text.replace("□全日制", "☑全日制")
            continue
        if label in label_map:
            v = label_map[label]
            new_t = f"{t.split('：')[0]}：{v}"
            if label == "论文题目":
                new_t = f"论文题目：{v}"
            for r in list(p.runs[1:]):
                r._element.getparent().remove(r._element)
            if p.runs:
                p.runs[0].text = new_t
            else:
                p.add_run(new_t)


def main():
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    OUT.mkdir(parents=True, exist_ok=True)
    shutil.copy(TEMPLATE, OUT_DOCX)
    os.chmod(OUT_DOCX, stat.S_IWRITE)
    doc = Document(str(OUT_DOCX))
    fill_cover(doc)

    draft = DRAFT.read_text(encoding="utf-8")
    lines = draft.splitlines()
    a = next(i for i, ln in enumerate(lines) if ln.startswith("**（1）研究背景"))
    b = next(i for i, ln in enumerate(lines) if ln.strip().startswith("### 4．阶段性研究成果"))
    blocks = build_progress_blocks(lines[a:b])

    t0 = doc.tables[0]
    for cell in t0.rows[0].cells:
        if "开题报告通过时间" in cell.text:
            for p in cell.paragraphs:
                for r in p.runs:
                    if "年" in r.text:
                        r.text = "1．开题报告通过时间：2025年12月24日（以开题记录为准）"
                break
            break
    for cell in t0.rows[2].cells:
        if "学分要求" in cell.text:
            for p in cell.paragraphs:
                for r in p.runs:
                    if "□是" in r.text:
                        r.text = r.text.replace("□是", "☑是")

    r4 = t0.rows[4].cells[0]
    instruction = "按照开题计划，填写开题以来学位论文工作的研究进展。视具体研究内容，可包括理论、计算、实验（或实证）等方面（可续页）"
    clear_cell(r4)
    instr_p = r4.add_paragraph()
    ir = instr_p.add_run(instruction)
    set_run(ir, size=12)
    for blk in blocks:
        if blk[0] == "h":
            add_para(r4, blk[1], bold=True)
        elif blk[0] == "p":
            add_para(r4, blk[1])
        elif blk[0] == "fig":
            add_figure(r4, blk[1], blk[2])
        elif blk[0] == "fig2":
            add_figure(r4, blk[1], "图7a HEADER_ONLY 操作端到端时延分布", 9.5)
            add_figure(r4, blk[2], "图7b BODY_ROTATION 操作端到端时延分布", 9.5)
        elif blk[0] == "table":
            add_table(r4, blk[1], blk[2])
        elif blk[0] == "algo":
            add_algo(r4, blk[1])

    r6 = t0.rows[6].cells[0]
    sres = ["按《研究生学位论文撰写格式规范》的格式要求分类填写与学位论文相关的阶段性研究成果（可续页）",
            "[1] 王威, 夏琦, 高建彬, 夏虎. 面向链上数据的双重绑定门限解封装方案[J]. 软件学报（拟投稿）.",
            "[2] [王威]. 一种非连续时间访问策略的压缩方法及系统: 中国, [P].（拟申请）.",
            "[3] [王威]. 一种基于属性与链上请求双重绑定的数据共享方法及系统: 中国, [P].（拟申请）.",
            "另：三套可复现原型与冻结实验数据集（时间策略编译、许可链授权执行、版本化密文头部与撤销恢复）。"]
    clear_cell(r6)
    for line in sres:
        add_para(r6, line, bold=line.startswith("["))

    c = draft.find("总体来看，当前研究已经形成")
    d = draft.find("针对上述问题，后续将围绕论文整合与理论深化")
    e = draft.find("## 三、中期考评审查意见")
    t1 = doc.tables[1]
    clear_cell(t1.rows[0].cells[0])
    add_para(t1.rows[0].cells[0], "1．未按开题计划完成的研究工作，研究工作存在的原理性、技术性难题以及在实验条件等方面的限制（可续页）", bold=True)
    for para in draft[c:d].splitlines():
        s = para.strip()
        if not s or s.startswith("## ") or s.startswith("### "):
            continue
        add_para(t1.rows[0].cells[0], s)
    clear_cell(t1.rows[1].cells[0])
    add_para(t1.rows[1].cells[0], "2．针对上述问题采取何种解决办法，对学位论文的研究内容及所采取的理论方法、技术路线和实施方案的进一步调整，以及下一步研究计划（可续页）", bold=True)
    for para in draft[d:e].splitlines():
        s = para.strip()
        if not s or s.startswith("## ") or s.startswith("### "):
            continue
        add_para(t1.rows[1].cells[0], s)

    doc.save(OUT_DOCX)
    print("saved:", OUT_DOCX)


if __name__ == "__main__":
    main()
