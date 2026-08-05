"""Dump the two official UESTC template DOCX files for I17."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


COVER = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\电子科技大学研究生学位论文封面及扉页 - 适用于专业学位硕士_081705087525.docx")
SPEC = Path(r"D:\Users\wangw\Documents\xwechat_files\wxid_qxnxx2moo0vz22_5966\msg\file\2026-08\电子科技大学研究生学位论文撰写规范- 适用于中国学生 - 副本_031543351520.docx")


def dump(path: Path, label: str, full: bool = True) -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print(f"\n################ {label} ################")
    doc = Document(str(path))
    print("sections:", len(doc.sections))
    for si, s in enumerate(doc.sections):
        print(f"  sec{si}: page {s.page_width.cm:.1f}x{s.page_height.cm:.1f}cm "
              f"margins T{s.top_margin.cm:.2f} B{s.bottom_margin.cm:.2f} L{s.left_margin.cm:.2f} R{s.right_margin.cm:.2f}")
    print("---- paragraphs ----")
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "?"
        txt = p.text
        if not txt.strip() and i > 0:
            continue
        fmt = []
        for r in p.runs[:2]:
            if r.font.size:
                fmt.append(f"sz={r.font.size.pt}")
            if r.font.name:
                fmt.append(f"font={r.font.name}")
            rpr = r._element.find(qn("w:rPr"))
            if rpr is not None:
                rf = rpr.find(qn("w:rFonts"))
                if rf is not None:
                    ea = rf.get(qn("w:eastAsia"))
                    if ea:
                        fmt.append(f"east={ea}")
            if r.bold:
                fmt.append("B")
            break
        print(f"[{i}] ({style}) {' '.join(fmt)} | {txt[:130]}")
    print("---- tables ----")
    for ti, t in enumerate(doc.tables):
        print(f"table {ti}: rows={len(t.rows)} cols={len(t.columns)}")
        for r in t.rows[:12]:
            print("   ", " || ".join(c.text.strip().replace("\n", " ")[:40] for c in r.cells)[:220])
    print("---- headers/footers ----")
    for si, s in enumerate(doc.sections):
        for kind, obj in (("header", s.header), ("footer", s.footer)):
            texts = [p.text for p in obj.paragraphs if p.text.strip()]
            if texts:
                print(f"  sec{si} {kind}: {texts}")
    # styles summary
    styles = {}
    for p in doc.paragraphs:
        if p.style is not None:
            styles[p.style.name] = styles.get(p.style.name, 0) + 1
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    if p.style is not None:
                        styles[p.style.name] = styles.get(p.style.name, 0) + 1
    print("---- style usage ----")
    for k, v in sorted(styles.items(), key=lambda kv: -kv[1]):
        print(f"  {k}: {v}")


def main() -> None:
    dump(COVER, "COVER & FLYLEAF TEMPLATE")
    dump(SPEC, "WRITING SPEC")


if __name__ == "__main__":
    main()
