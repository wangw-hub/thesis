# -*- coding: utf-8 -*-
"""FINAL-CLEAN structural QA against the built DOCX and FINAL source."""
from __future__ import annotations

import json
import re
import sys

from docx import Document
from lxml import etree


ROOT = r"D:\Research\crypto_thesis\epoch-authorization-r3-prep"
DOCX = ROOT + r"\docs\midterm-report\final\output\王威-专业学位研究生学位论文中期考评表-最终固化版.docx"
SRC = ROOT + r"\docs\midterm-report\final\FINAL-MIDTERM-SOURCE.md"
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
M = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    report: dict = {}

    src = open(SRC, encoding="utf-8").read()
    src_body = src[: src.find("### 参考文献")]

    # citation order
    first_seen: dict[int, int] = {}
    for m in re.finditer(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]", src_body):
        nums = []
        for part in m.group(0)[1:-1].replace(" ", "").split(","):
            if "-" in part:
                a, b = part.split("-")
                nums.extend(range(int(a), int(b) + 1))
            else:
                nums.append(int(part))
        for n in nums:
            if n not in first_seen:
                first_seen[n] = m.start()
    seq = sorted(first_seen.items(), key=lambda kv: kv[1])
    nums = [n for n, _ in seq]
    report["citation_count"] = len(nums)
    report["citation_order_error"] = nums != list(range(1, len(nums) + 1))

    # references
    ref_block = src[src.find("### 参考文献"): src.find("### 4．阶段性研究成果")]
    refs = re.findall(r"^\[(\d+)\]", ref_block, re.M)
    report["reference_count"] = len(refs)
    report["reference_numbers_ok"] = [int(x) for x in refs] == list(range(1, 35))
    report["access_date_visible"] = ref_block.count("[2026-08-02]") + ref_block.count("2014[")
    entries = re.split(r"\n\n\[", ref_block)
    y2021 = y2024 = 0
    for e in entries:
        ym = re.findall(r"(20\d\d)", e)
        if not ym:
            continue
        y = int(ym[-1])
        if y >= 2021:
            y2021 += 1
        if y >= 2024:
            y2024 += 1
    report["ref_2021_2026"] = y2021
    report["ref_2024_2026"] = y2024
    report["ref_ratio"] = round(y2021 / len(refs), 3)

    cited = []
    for m in re.finditer(r"\[\d+(?:-\d+)?(?:,\s*\d+(?:-\d+)?)*\]", src_body):
        for part in m.group(0)[1:-1].replace(" ", "").split(","):
            if "-" in part:
                a, b = part.split("-")
                cited.extend(range(int(a), int(b) + 1))
            else:
                cited.append(int(part))
    cited_set = set(cited)
    report["orphan_references"] = [n for n in range(1, len(refs) + 1) if n not in cited_set]

    # algorithms
    algos = re.findall(r"\[算法块：算法(\d+) ([^\n]+)", src)
    report["algorithms"] = [f"{a} {b}" for a, b in algos]
    report["algo_numbering_ok"] = [int(a) for a, _ in algos] == list(range(1, 9))
    report["formula_count"] = len(re.findall(r"^\[公式：", src, re.M))

    # symbols
    sym_checks = {
        "U_user_tuple": r"五元组 \(U_u=(account",
        "R_d_redundancy": r"冗余度 \(R_d=2\)",
        "B_cap_formula": r"B_{cap}=\operatorname{Encode}",
        "B_cap_sig": r"sk_I,B_{cap}",
    }
    report["symbols"] = {k: (v in src) for k, v in sym_checks.items()}

    # DOCX checks
    doc = Document(DOCX)
    body = doc.element.body
    empty_nary = 0
    for om in body.iter(M + "oMath"):
        if "<m:e/>" in etree.tostring(om, encoding="unicode"):
            empty_nary += 1
    report["docx_empty_nary"] = empty_nary

    all_text = "\n".join(p.text or "" for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                all_text += "\n" + (cell.text or "")
    report["algo_end_text_in_docx"] = "算法结束" in all_text

    tables = doc.tables
    nested = []
    for tbl in tables:
        nested.extend(tbl._tbl.findall(f".//{W}tbl"))
    all_tbls = [t._tbl for t in tables] + nested
    report["docx_tables"] = len(all_tbls)
    missing_hdr = 0
    for tbl in all_tbls:
        hdr = tbl.findall(f".//{W}trPr/{W}tblHeader")
        if not hdr:
            missing_hdr += 1
    report["tables_without_repeat_header"] = missing_hdr

    check_chars = 0
    sym_count = 0
    for para in doc.paragraphs:
        if "攻读学位级别" in (para.text or ""):
            check_chars = (para.text or "").count("☑")
            sym_count = len(para._p.findall(f".//{W}sym"))
    report["cover_checked_marks"] = check_chars
    report["cover_wingdings_sym"] = sym_count

    figs = len(body.findall(f".//{W}drawing")) + len(body.findall(f".//{W}pict"))
    report["docx_figures"] = figs

    print(json.dumps(report, ensure_ascii=False, indent=1))


if __name__ == "__main__":
    main()
